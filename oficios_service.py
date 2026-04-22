from __future__ import annotations

import argparse
import base64
import hashlib
import json
import logging
import os
import re
import sqlite3
import sys
import threading
import time
import tkinter as tk
from collections import Counter
from dataclasses import dataclass, field
from datetime import datetime, date, timedelta
from pathlib import Path
from tkinter import messagebox, ttk
from typing import Any, Callable, Dict, List, Optional, Tuple
from zoneinfo import ZoneInfo

try:
    import customtkinter as ctk
    CTK_AVAILABLE = True
except ImportError:
    CTK_AVAILABLE = False

try:
    from docx import Document as DocxDocument
    from docx.shared import Pt, Inches, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False

import msal
import requests
from openpyxl import Workbook, load_workbook
from openpyxl.styles import Font, PatternFill, Alignment, Border, Side

EXPECTED_COLUMNS = [
    "Nro",
    "Categoría",
    "Fecha de Oficio",
    "Concepto",
    "Dirección Responsable",
    "Gerencia Responsable",
    "Gerente Responsable",
    "Equipo",
    "Plazo Respuesta",
    "Multa",
]

AREAS_VALIDAS = [
    "PMGD", "Conexiones", "Lectura",
    "Servicio al Cliente", "Cobranza", "Pérdidas",
]

SCHEMA = {
    "type": "object",
    "additionalProperties": False,
    "properties": {
        "numero_oficio": {"type": ["string", "null"]},
        "categoria": {
            "type": ["string", "null"],
            "enum": ["Resolución exenta", "Oficio ordinario", "Oficio circular", None],
        },
        "fecha_oficio": {"type": ["string", "null"]},
        "concepto": {"type": ["string", "null"]},
        "gerencia_responsable": {
            "type": ["string", "null"],
            "enum": ["PMGD", "Conexiones", "Lectura", "Servicio al Cliente", "Cobranza", "Pérdidas", None],
        },
        "plazo_respuesta": {"type": ["string", "null"]},
        "plazo_relativo_cantidad": {"type": ["integer", "null"], "minimum": 1},
        "plazo_relativo_tipo": {
            "type": ["string", "null"],
            "enum": ["dias_corridos", "dias_habiles", None],
        },
        "oficio_relacionado": {"type": ["string", "null"]},
        "remitente": {"type": ["string", "null"]},
        "keywords": {
            "type": "array",
            "items": {"type": "string"},
        },
        "confianza": {"type": ["number", "null"], "minimum": 0, "maximum": 1},
    },
    "required": [
        "numero_oficio",
        "categoria",
        "fecha_oficio",
        "concepto",
        "gerencia_responsable",
        "plazo_respuesta",
        "plazo_relativo_cantidad",
        "plazo_relativo_tipo",
        "oficio_relacionado",
        "remitente",
        "keywords",
        "confianza",
    ],
}

PROMPT = """Eres un extractor de metadatos regulatorios para oficios y resoluciones recibidos por CGE.

Tu tarea es leer el PDF adjunto y devolver SOLO un JSON válido que cumpla exactamente el schema entregado.

Reglas de extracción y normalización:

1. Usa exclusivamente información presente en el PDF.
2. numero_oficio:
   - devuelve solo el número principal del oficio o resolución
   - sin \"N°\", sin \"Nº\", sin \"No.\", sin prefijos
   - ejemplo: \"Resolución Exenta Electrónica N° 38222\" => \"38222\"
3. categoria:
   normaliza a uno de estos tres valores exactos según el prefijo del nombre del archivo:
   - \"RE\" => \"Resolución exenta\"
   - \"Ord.\" => \"Oficio ordinario\"
   - \"OC\" => \"Oficio circular\"
   Usa siempre el prefijo del nombre del archivo para determinar la categoría.
4. fecha_oficio:
   - es la fecha de emisión/envío del oficio o resolución
   - formato exacto YYYY-MM-DD
   - si no está clara, devuelve null
5. concepto:
   - resume el asunto principal del documento
   - una sola línea
   - máximo 250 caracteres
   - sin comillas, sin saltos de línea
   - debe ser ejecutivo y fiel al contenido
6. gerencia_responsable:
   normaliza a uno de estos valores exactos:
   - \"PMGD\"
   - \"Conexiones\"
   - \"Lectura\"
   - \"Servicio al Cliente\"
   - \"Cobranza\"
   - \"Pérdidas\"
   Criterios:
   - \"PMGD\": generación distribuida, conexión de PMGD, plataformas o procesos PMGD
   - \"Conexiones\": conexión, empalme, factibilidad, puesta en servicio, plazos de conexión, obras o procesos de conexión
   - \"Lectura\": lectura de medidores, medición, consumos leídos, toma de lectura
   - \"Servicio al Cliente\": atención, reclamos, canales, calidad de servicio, respuesta al cliente
   - \"Cobranza\": deuda, mora, pago, repactación, suspensión/corte por deuda, cobranza
   - \"Pérdidas\": pérdidas de energía, hurto de energía, consumo no registrado, fraude eléctrico, intervención de medidor, irregularidades en consumo
   Si hay más de un tema, elige el principal.
7. plazo_respuesta:
   - devuelve una fecha explícita de vencimiento en formato YYYY-MM-DD solo si en el PDF existe una fecha calendario inequívoca
   - si no existe una fecha exacta, devuelve null
8. plazo_relativo_cantidad y plazo_relativo_tipo:
   - si el documento establece un plazo relativo, extráelo
   - ejemplos:
     - \"10 días hábiles\" => plazo_relativo_cantidad=10, plazo_relativo_tipo=\"dias_habiles\"
     - \"5 días corridos\" => plazo_relativo_cantidad=5, plazo_relativo_tipo=\"dias_corridos\"
     - \"30 días\" sin aclaración expresa => trátalo como \"dias_corridos\"
   - si existe una fecha exacta y además un plazo relativo, devuelve ambos si ambos aparecen explícitamente
   - si no existe plazo relativo, devuelve null en ambos campos
9. oficio_relacionado:
   - si el documento hace referencia explícita a otro oficio, resolución u orden de compra anterior
     (expresiones como "en respuesta a", "en relación al Oficio N°", "complementa el Ord.", "adjunto al OC"),
     devuelve el número de ese documento (solo cifras o código, sin prefijos)
   - si no hay referencia a otro documento, devuelve null
10. remitente:
   - nombre de la persona u organismo que firma/envía el oficio (ej: "Superintendente SEC", "Jefe División Fiscalización")
   - si no está claro, devuelve null
11. keywords:
   - lista de 3 a 6 palabras o frases cortas del documento que justifican la clasificación del área
   - deben ser términos presentes en el PDF (no inventados)
12. confianza:
   - número entre 0.0 y 1.0 que refleja qué tan seguro estás de la clasificación del área (gerencia_responsable)
   - 1.0 = el PDF menciona explícitamente el área o sus conceptos; 0.5 = dos áreas compiten; <0.5 = clasificación incierta
13. Si un dato no está claro o no existe, devuelve null.
14. No agregues ningún texto fuera del JSON.
15. No inventes datos.
"""

INFORME_MULTA_SCHEMA = {
    "type": "object",
    "additionalProperties": False,
    "properties": {
        "sociedad":                 {"type": ["string", "null"]},
        "regional":                 {"type": ["string", "null"]},
        "zonal":                    {"type": ["string", "null"]},
        "evento":                   {"type": ["string", "null"]},
        "sector_activo":            {"type": ["string", "null"]},
        "clientes_afectados":       {"type": ["string", "null"]},
        "fecha_evento":             {"type": ["string", "null"]},
        "fecha_notificacion_multa": {"type": ["string", "null"]},
        "monto_utm":                {"type": ["string", "null"]},
        "numero_resolucion":        {"type": ["string", "null"]},
        "etapa_multa":              {"type": ["string", "null"]},
        "descripcion_motivo_sec":   {"type": ["string", "null"]},
        "descripcion_tecnica_zona": {"type": ["string", "null"]},
        "causa_raiz":               {"type": ["string", "null"]},
        "cronologia":               {"type": ["string", "null"]},
        "que_paso":                 {"type": ["string", "null"]},
        "que_se_hizo":              {"type": ["string", "null"]},
        "plan_mejora":              {"type": ["string", "null"]},
        "area_responsable_multa":   {"type": ["string", "null"]},
        "area_responsable_mejora":  {"type": ["string", "null"]},
    },
    "required": [
        "sociedad", "regional", "zonal", "evento", "sector_activo",
        "clientes_afectados", "fecha_evento", "fecha_notificacion_multa",
        "monto_utm", "numero_resolucion", "etapa_multa",
        "descripcion_motivo_sec", "descripcion_tecnica_zona", "causa_raiz",
        "cronologia", "que_paso", "que_se_hizo", "plan_mejora",
        "area_responsable_multa", "area_responsable_mejora",
    ],
}

INFORME_MULTA_PROMPT = """Eres un analista regulatorio de CGE. A partir del PDF adjunto (una multa o resolución sancionatoria de la SEC), extrae SOLO un JSON válido con los campos del informe de zona.

Instrucciones:
1. sociedad: nombre de la empresa sancionada (ej: "CGE Distribución S.A.").
2. regional: región geográfica involucrada (ej: "Metropolitana", "Tarapacá").
3. zonal: zona operativa específica si se menciona.
4. evento: descripción breve del tipo de evento sancionado.
5. sector_activo: sector eléctrico o activo físico involucrado.
6. clientes_afectados: número o descripción de clientes afectados, o null.
7. fecha_evento: fecha del evento sancionado (formato DD-MM-YYYY o texto si no es exacta).
8. fecha_notificacion_multa: fecha en que CGE fue notificada de la multa.
9. monto_utm: monto de la multa en UTM como texto (ej: "100 UTM").
10. numero_resolucion: número de la resolución exenta SEC.
11. etapa_multa: etapa actual (ej: "Notificación", "Descargos", "Resolución firme").
12. descripcion_motivo_sec: descripción objetiva según la SEC del motivo de la multa. Sé detallado.
13. descripcion_tecnica_zona: descripción técnica del evento desde la perspectiva operativa.
14. causa_raiz: posible causa raíz interna identificable en el documento.
15. cronologia: secuencia cronológica de los hechos relevantes.
16. que_paso: resumen de qué ocurrió y cuál es la problemática sancionada.
17. que_se_hizo: qué acciones se tomaron o dejaron de tomar.
18. plan_mejora: si existe algún plan de mejora mencionado, descríbelo. Si no, null.
19. area_responsable_multa: área interna de CGE responsable de la infracción.
20. area_responsable_mejora: área responsable de implementar correcciones.

Si un dato no está en el PDF, devuelve null. No inventes datos. Solo JSON, sin texto adicional.
"""


def get_base_dir() -> Path:
    """Retorna el directorio base del ejecutable o del script."""
    if getattr(sys, "frozen", False):
        return Path(sys.executable).parent
    return Path(__file__).parent


@dataclass
class Gerente:
    nombre: str
    email: str = ""


@dataclass
class PlannerConfig:
    enabled: bool = False
    tenant_id: str = ""
    client_id: str = ""
    client_secret: str = ""
    plan_id: str = ""
    bucket_id: str = ""


@dataclass
class OutlookConfig:
    enabled: bool = False
    user_email: str = ""
    folder_name: str = "Oficios"


@dataclass
class Config:
    watch_dir: Path
    excel_path: Path
    processed_state_path: Path
    corrections_path: Path
    historial_path: Path
    reglas_path: Path
    log_path: Path
    timezone: str
    run_time: str
    openai_api_key: str
    model: str
    gerentes: Dict[str, Gerente]
    planner: PlannerConfig = field(default_factory=PlannerConfig)
    informe_multa_api_key: str = ""
    informe_multa_model: str = "claude-sonnet-4-20250514"
    db_path: Path = field(default_factory=lambda: Path("oficios.db"))
    informe_output_dir: Path = field(default_factory=lambda: Path("."))
    request_timeout_seconds: int = 180
    scan_extensions: tuple[str, ...] = (".pdf",)


def load_config(path: Path) -> Config:
    with path.open("r", encoding="utf-8") as f:
        try:
            raw = json.load(f)
        except json.JSONDecodeError as exc:
            raise ValueError(
                f"El archivo de configuración tiene un error de formato JSON:\n"
                f"  Archivo: {path}\n"
                f"  Detalle: {exc}\n\n"
                f"Revise que no haya comas sobrantes, comillas sin cerrar o caracteres inválidos."
            ) from exc

    gerentes: Dict[str, Gerente] = {}
    for area, item in raw.get("gerentes", {}).items():
        gerentes[area] = Gerente(nombre=item.get("nombre", ""), email=item.get("email", ""))

    api_key = raw.get("openai_api_key") or os.getenv("OPENAI_API_KEY", "")
    if not api_key:
        raise ValueError("Falta openai_api_key en config.json o variable de entorno OPENAI_API_KEY.")

    planner_raw = raw.get("planner", {})
    planner = PlannerConfig(
        enabled=planner_raw.get("enabled", False),
        tenant_id=planner_raw.get("tenant_id", ""),
        client_id=planner_raw.get("client_id", ""),
        client_secret=planner_raw.get("client_secret", ""),
        plan_id=planner_raw.get("plan_id", ""),
        bucket_id=planner_raw.get("bucket_id", ""),
    )

    informe_raw = raw.get("informe_multa", {})
    excel_parent = str(Path(raw["excel_path"]).parent)
    informe_output_dir = Path(informe_raw.get("output_dir", excel_parent))

    excel_p = Path(raw["excel_path"])
    db_p = Path(raw.get("db_path", str(excel_p.with_suffix(".db"))))

    return Config(
        watch_dir=Path(raw["watch_dir"]),
        excel_path=excel_p,
        processed_state_path=Path(raw.get("processed_state_path", "processed_state.json")),
        corrections_path=Path(raw.get("corrections_path", "corrections.json")),
        historial_path=Path(raw.get("historial_path", "historial_oficios.json")),
        reglas_path=Path(raw.get("reglas_path", "reglas_clasificacion.json")),
        log_path=Path(raw.get("log_path", "oficios_service.log")),
        timezone=raw.get("timezone", "America/Santiago"),
        run_time=raw.get("run_time", "16:00"),
        openai_api_key=api_key,
        model=raw.get("model", "gpt-5.4-mini"),
        gerentes=gerentes,
        db_path=db_p,
        planner=planner,
        informe_multa_api_key=informe_raw.get("api_key", ""),
        informe_multa_model=informe_raw.get("model", "claude-sonnet-4-20250514"),
        informe_output_dir=informe_output_dir,
        request_timeout_seconds=int(raw.get("request_timeout_seconds", 180)),
    )


def setup_logging(log_path: Path) -> None:
    log_path.parent.mkdir(parents=True, exist_ok=True)
    logging.basicConfig(
        level=logging.INFO,
        format="%(asctime)s | %(levelname)s | %(message)s",
        handlers=[
            logging.FileHandler(log_path, encoding="utf-8"),
            logging.StreamHandler(sys.stdout),
        ],
    )


def ensure_parent(path: Path) -> None:
    path.parent.mkdir(parents=True, exist_ok=True)


def create_excel_template(path: Path, sheet_name: str = "Oficios") -> None:
    ensure_parent(path)
    wb = Workbook()
    ws = wb.active
    ws.title = sheet_name
    ws.freeze_panes = "A2"

    header_fill = PatternFill(fill_type="solid", fgColor="1F4E78")
    header_font = Font(color="FFFFFF", bold=True)
    thin_gray = Side(style="thin", color="D9E2F3")

    for col_idx, name in enumerate(EXPECTED_COLUMNS, start=1):
        cell = ws.cell(row=1, column=col_idx, value=name)
        cell.fill = header_fill
        cell.font = header_font
        cell.alignment = Alignment(horizontal="center", vertical="center")
        cell.border = Border(bottom=thin_gray)

    widths = {
        "A": 12,
        "B": 22,
        "C": 16,
        "D": 60,
        "E": 28,
        "F": 24,
        "G": 26,
        "H": 14,
        "I": 18,
        "J": 10,
    }
    for col, width in widths.items():
        ws.column_dimensions[col].width = width

    for col in ["C", "I"]:
        for row in range(2, 5000):
            ws[f"{col}{row}"].number_format = "DD-MM-YYYY"

    wb.save(path)


def ensure_excel_exists(path: Path) -> None:
    if not path.exists():
        create_excel_template(path)
        logging.info("Plantilla Excel creada en %s", path)
        return

    wb = load_workbook(path)
    ws = wb.active
    existing = [ws.cell(row=1, column=i).value for i in range(1, len(EXPECTED_COLUMNS) + 1)]
    if existing == EXPECTED_COLUMNS:
        return

    # Migrar: si solo falta la columna "Multa" al final, agregarla automáticamente
    old_columns = EXPECTED_COLUMNS[:-1]
    existing_old = [ws.cell(row=1, column=i).value for i in range(1, len(old_columns) + 1)]
    if existing_old == old_columns:
        col_idx = len(EXPECTED_COLUMNS)
        header_cell = ws.cell(row=1, column=col_idx, value="Multa")
        header_cell.fill = PatternFill(fill_type="solid", fgColor="1F4E78")
        header_cell.font = Font(color="FFFFFF", bold=True)
        header_cell.alignment = Alignment(horizontal="center", vertical="center")
        header_cell.border = Border(bottom=Side(style="thin", color="D9E2F3"))
        ws.column_dimensions["J"].width = 10
        # Rellenar multa para filas existentes según concepto
        for row_idx in range(2, ws.max_row + 1):
            concepto = str(ws.cell(row=row_idx, column=4).value or "")
            if concepto and es_multa_o_cargos(concepto):
                cell = ws.cell(row=row_idx, column=col_idx, value="Sí")
                cell.alignment = Alignment(horizontal="center", vertical="center")
        wb.save(path)
        logging.info("Columna 'Multa' agregada al Excel existente.")
        return

    raise ValueError(
        f"El Excel existente no tiene los encabezados esperados. Esperado: {EXPECTED_COLUMNS} | Actual: {existing}"
    )


def load_state(path: Path) -> dict[str, Any]:
    if not path.exists():
        return {"processed_hashes": [], "last_run_date": None}
    with path.open("r", encoding="utf-8") as f:
        return json.load(f)


def save_state(path: Path, state: dict[str, Any]) -> None:
    ensure_parent(path)
    with path.open("w", encoding="utf-8") as f:
        json.dump(state, f, ensure_ascii=False, indent=2)


def sha256_file(path: Path) -> str:
    h = hashlib.sha256()
    with path.open("rb") as f:
        for chunk in iter(lambda: f.read(1024 * 1024), b""):
            h.update(chunk)
    return h.hexdigest()


def parse_date_yyyy_mm_dd(value: Optional[str]) -> Optional[date]:
    if not value:
        return None
    try:
        return datetime.strptime(value, "%Y-%m-%d").date()
    except ValueError:
        return None


def add_business_days(start_date: date, business_days: int) -> date:
    current = start_date
    added = 0
    while added < business_days:
        current += timedelta(days=1)
        if current.weekday() < 5:  # 0-4 = lunes-viernes
            added += 1
    return current


def compute_due_date(extracted: dict[str, Any]) -> Optional[date]:
    explicit_due = parse_date_yyyy_mm_dd(extracted.get("plazo_respuesta"))
    if explicit_due:
        return explicit_due

    fecha_oficio = parse_date_yyyy_mm_dd(extracted.get("fecha_oficio"))
    if not fecha_oficio:
        return None

    cantidad = extracted.get("plazo_relativo_cantidad")
    tipo = extracted.get("plazo_relativo_tipo")

    if not isinstance(cantidad, int) or cantidad <= 0 or tipo not in {"dias_corridos", "dias_habiles"}:
        return None

    if tipo == "dias_habiles":
        return add_business_days(fecha_oficio, cantidad)

    return fecha_oficio + timedelta(days=cantidad)


def extract_output_text(api_response: dict[str, Any]) -> str:
    if isinstance(api_response.get("output_text"), str):
        return api_response["output_text"]

    pieces: List[str] = []
    for item in api_response.get("output", []):
        for content in item.get("content", []):
            if content.get("type") in {"output_text", "text"} and isinstance(content.get("text"), str):
                pieces.append(content["text"])
    if pieces:
        return "\n".join(pieces).strip()

    raise ValueError("No se pudo extraer output_text de la respuesta de OpenAI.")


def call_openai_extract(
    config: Config,
    pdf_path: Path,
    corrections_prompt: str = "",
    related_pdf_path: Optional[Path] = None,
) -> dict[str, Any]:
    pdf_b64 = base64.b64encode(pdf_path.read_bytes()).decode("ascii")
    full_prompt = PROMPT + corrections_prompt

    if related_pdf_path is not None:
        intro_text = (
            f"Se adjuntan DOS documentos vinculados entre sí. "
            f"Documento principal: {pdf_path.name}. "
            f"Documento relacionado: {related_pdf_path.name}. "
            f"Estudia ambos para determinar el área responsable y los demás metadatos. "
            f"El JSON debe corresponder al documento principal."
        )
        related_b64 = base64.b64encode(related_pdf_path.read_bytes()).decode("ascii")
        user_content = [
            {"type": "input_text", "text": intro_text},
            {"type": "input_file", "filename": pdf_path.name,
             "file_data": f"data:application/pdf;base64,{pdf_b64}"},
            {"type": "input_file", "filename": related_pdf_path.name,
             "file_data": f"data:application/pdf;base64,{related_b64}"},
        ]
    else:
        user_content = [
            {"type": "input_text",
             "text": f"Analiza el siguiente PDF. Nombre del archivo: {pdf_path.name}"},
            {"type": "input_file", "filename": pdf_path.name,
             "file_data": f"data:application/pdf;base64,{pdf_b64}"},
        ]

    payload = {
        "model": config.model,
        "max_output_tokens": 450,
        "text": {
            "format": {
                "type": "json_schema",
                "name": "oficio_sec_extract",
                "strict": True,
                "schema": SCHEMA,
            }
        },
        "input": [
            {
                "role": "developer",
                "content": [{"type": "input_text", "text": full_prompt}],
            },
            {
                "role": "user",
                "content": user_content,
            },
        ],
    }

    headers = {
        "Content-Type": "application/json",
        "Authorization": f"Bearer {config.openai_api_key}",
    }

    response = requests.post(
    "https://api.openai.com/v1/responses",
    headers=headers,
    json=payload,
    timeout=config.request_timeout_seconds,
    )

    if not response.ok:
        logging.error("OpenAI devolvió %s: %s", response.status_code, response.text)
        response.raise_for_status()

    data = response.json()

    output_text = extract_output_text(data)
    try:
        parsed = json.loads(output_text)
    except json.JSONDecodeError as exc:
        raise ValueError(f"La respuesta de OpenAI no es JSON válido: {output_text}") from exc

    usage = data.get("usage", {})
    if usage:
        logging.info(
            "Uso OpenAI para %s | input_tokens=%s | output_tokens=%s | total_tokens=%s",
            pdf_path.name,
            usage.get("input_tokens"),
            usage.get("output_tokens"),
            usage.get("total_tokens"),
        )

    return parsed


def map_row(extracted: dict[str, Any], gerentes: Dict[str, Gerente]) -> List[Any]:
    gerencia = extracted.get("gerencia_responsable")
    gerente = gerentes.get(gerencia, Gerente(nombre=""))
    fecha_oficio = parse_date_yyyy_mm_dd(extracted.get("fecha_oficio"))
    plazo = compute_due_date(extracted)

    concepto = extracted.get("concepto") or ""
    multa = "Sí" if es_multa_o_cargos(concepto) else ""

    return [
        extracted.get("numero_oficio") or "",
        extracted.get("categoria") or "",
        fecha_oficio,
        concepto,
        "Comercial y Servicio al Cliente",
        gerencia or "",
        gerente.nombre,
        "",
        plazo,
        multa,
    ]


def row_exists(ws, nro: str, categoria: str, fecha_oficio: Optional[date]) -> bool:
    target_date_str = fecha_oficio.isoformat() if fecha_oficio else ""
    for row in ws.iter_rows(min_row=2, values_only=True):
        existing_nro = str(row[0] or "").strip()
        existing_cat = str(row[1] or "").strip()
        existing_fecha = row[2]
        if hasattr(existing_fecha, "date"):
            existing_fecha = existing_fecha.date()
        existing_fecha_str = existing_fecha.isoformat() if existing_fecha else ""
        if existing_nro == nro and existing_cat == categoria and existing_fecha_str == target_date_str:
            return True
    return False


def first_empty_row(ws) -> int:
    for row_idx in range(2, ws.max_row + 1):
        row_values = [ws.cell(row=row_idx, column=col_idx).value for col_idx in range(1, len(EXPECTED_COLUMNS) + 1)]
        if all(v is None or str(v).strip() == "" for v in row_values):
            return row_idx
    return ws.max_row + 1


def append_to_excel(excel_path: Path, row: List[Any]) -> bool:
    wb = load_workbook(excel_path)
    ws = wb.active

    nro = str(row[0] or "").strip()
    categoria = str(row[1] or "").strip()
    fecha_oficio = row[2] if isinstance(row[2], date) else None

    if row_exists(ws, nro, categoria, fecha_oficio):
        logging.info("Se omite inserción duplicada en Excel: nro=%s | categoría=%s | fecha=%s", nro, categoria, fecha_oficio)
        return False

    next_row = first_empty_row(ws)

    for col_idx, value in enumerate(row, start=1):
        cell = ws.cell(row=next_row, column=col_idx, value=value)
        if col_idx in (3, 9):
            cell.number_format = "DD-MM-YYYY"
            cell.alignment = Alignment(horizontal="center")
        elif col_idx == 10:
            cell.alignment = Alignment(horizontal="center", vertical="center")
        elif col_idx in (1, 2, 5, 6, 7, 8):
            cell.alignment = Alignment(vertical="center")
        else:
            cell.alignment = Alignment(wrap_text=True, vertical="top")

    wb.save(excel_path)
    return True


_COPY_PATTERN = re.compile(
    r"^(?P<base>.+?)"
    r"(?:"
    r"\s*-\s*(?:cop(?:y|ia))(?:\s*\(\d+\))?"  # " - Copy", " - copia", " - Copy (2)"
    r"|\s+\(\d+\)"                              # " (1)", " (2)"
    r")"
    r"(?P<ext>\.[^.]+)$",
    re.IGNORECASE,
)


def remove_duplicate_files(watch_dir: Path, extensions: tuple[str, ...]) -> int:
    """Detecta archivos que son copias (por nombre) y los elimina si el original existe."""
    removed = 0
    for path in sorted(watch_dir.iterdir()):
        if not path.is_file() or path.suffix.lower() not in extensions:
            continue
        m = _COPY_PATTERN.match(path.name)
        if not m:
            continue
        original = watch_dir / (m.group("base") + m.group("ext"))
        if original.exists() and original != path:
            logging.info("Archivo duplicado detectado: %s (original: %s). Eliminando copia.", path.name, original.name)
            path.unlink()
            removed += 1
    if removed:
        logging.info("Se eliminaron %d copia(s) de archivos.", removed)
    return removed


# ---------------------------------------------------------------------------
# SQLite storage (primary data store; Excel kept as export)
# ---------------------------------------------------------------------------

_DB_SCHEMA = """
CREATE TABLE IF NOT EXISTS oficios (
    id INTEGER PRIMARY KEY AUTOINCREMENT,
    nro TEXT NOT NULL,
    categoria TEXT DEFAULT '',
    fecha_oficio TEXT,
    concepto TEXT DEFAULT '',
    direccion TEXT DEFAULT 'Comercial y Servicio al Cliente',
    gerencia TEXT DEFAULT '',
    gerente TEXT DEFAULT '',
    equipo TEXT DEFAULT '',
    plazo_respuesta TEXT,
    multa TEXT DEFAULT '',
    created_at TEXT DEFAULT (datetime('now')),
    UNIQUE(nro, categoria, fecha_oficio)
);
CREATE INDEX IF NOT EXISTS idx_oficios_nro ON oficios(nro);
CREATE INDEX IF NOT EXISTS idx_oficios_plazo ON oficios(plazo_respuesta);
CREATE INDEX IF NOT EXISTS idx_oficios_gerencia ON oficios(gerencia);
"""


def init_db(db_path: Path) -> None:
    ensure_parent(db_path)
    con = sqlite3.connect(db_path)
    con.executescript(_DB_SCHEMA)
    con.close()


def migrate_excel_to_db(excel_path: Path, db_path: Path) -> int:
    """Import existing Excel rows into SQLite (skips duplicates). Returns count."""
    if not excel_path.exists():
        return 0
    con = sqlite3.connect(db_path)
    cur = con.cursor()
    cur.execute("SELECT COUNT(*) FROM oficios")
    if cur.fetchone()[0] > 0:
        con.close()
        return 0

    imported = 0
    try:
        wb = load_workbook(excel_path, read_only=True, data_only=True)
        ws = wb.active
        for row in ws.iter_rows(min_row=2, values_only=True):
            nro = str(row[0] or "").strip()
            if not nro:
                continue
            fecha = row[2]
            if hasattr(fecha, "date"):
                fecha = fecha.date()
            fecha_str = fecha.isoformat() if isinstance(fecha, date) else None
            plazo = row[8]
            if hasattr(plazo, "date"):
                plazo = plazo.date()
            plazo_str = plazo.isoformat() if isinstance(plazo, date) else None
            try:
                cur.execute(
                    "INSERT OR IGNORE INTO oficios "
                    "(nro,categoria,fecha_oficio,concepto,direccion,gerencia,gerente,equipo,plazo_respuesta,multa) "
                    "VALUES (?,?,?,?,?,?,?,?,?,?)",
                    (
                        nro,
                        str(row[1] or ""),
                        fecha_str,
                        str(row[3] or ""),
                        str(row[4] or "Comercial y Servicio al Cliente"),
                        str(row[5] or ""),
                        str(row[6] or ""),
                        str(row[7] or ""),
                        plazo_str,
                        str(row[9] or ""),
                    ),
                )
                imported += cur.rowcount
            except sqlite3.Error:
                pass
        wb.close()
        con.commit()
    except Exception as exc:
        logging.warning("Error migrando Excel a SQLite: %s", exc)
    finally:
        con.close()
    if imported:
        logging.info("Migrados %d oficios del Excel a SQLite.", imported)
    return imported


def db_insert_oficio(db_path: Path, row: List[Any]) -> bool:
    """Insert a row into SQLite. Returns True if inserted (not duplicate)."""
    fecha = row[2]
    fecha_str = fecha.isoformat() if isinstance(fecha, date) else None
    plazo = row[8]
    plazo_str = plazo.isoformat() if isinstance(plazo, date) else None
    con = sqlite3.connect(db_path)
    try:
        con.execute(
            "INSERT OR IGNORE INTO oficios "
            "(nro,categoria,fecha_oficio,concepto,direccion,gerencia,gerente,equipo,plazo_respuesta,multa) "
            "VALUES (?,?,?,?,?,?,?,?,?,?)",
            (
                str(row[0] or ""),
                str(row[1] or ""),
                fecha_str,
                str(row[3] or ""),
                str(row[4] or "Comercial y Servicio al Cliente"),
                str(row[5] or ""),
                str(row[6] or ""),
                str(row[7] or ""),
                plazo_str,
                str(row[9] or ""),
            ),
        )
        con.commit()
        return con.total_changes > 0
    finally:
        con.close()


def db_update_oficio(db_path: Path, nro: str, categoria: str,
                     updates: Dict[str, Any], gerentes: Dict[str, "Gerente"]) -> None:
    """Update fields in SQLite for a given oficio."""
    sets: List[str] = []
    vals: List[Any] = []
    if "gerencia_responsable" in updates:
        new_area = updates["gerencia_responsable"]
        sets.append("gerencia=?")
        vals.append(new_area)
        gerente = gerentes.get(new_area, Gerente(nombre=""))
        sets.append("gerente=?")
        vals.append(gerente.nombre)
    if "plazo_respuesta" in updates:
        p = parse_date_yyyy_mm_dd(updates["plazo_respuesta"])
        sets.append("plazo_respuesta=?")
        vals.append(p.isoformat() if p else None)
    if "es_multa" in updates:
        sets.append("multa=?")
        vals.append("Sí" if updates["es_multa"] else "")
    if not sets:
        return
    vals.extend([nro, categoria])
    sql = f"UPDATE oficios SET {','.join(sets)} WHERE nro=? AND categoria=?"
    con = sqlite3.connect(db_path)
    try:
        con.execute(sql, vals)
        con.commit()
    finally:
        con.close()


_VALID_PREFIXES = re.compile(r"^(OC|Ord\.|RE)\s", re.IGNORECASE)


def find_related_pdf(watch_dir: Path, related_nro: str) -> Optional[Path]:
    """Busca en watch_dir un PDF cuyo nombre contenga el número de oficio relacionado."""
    if not related_nro or not watch_dir.exists():
        return None
    # Normalizar: quitar espacios y guiones para comparación flexible
    nro_clean = re.sub(r"[\s\-/]", "", related_nro).lower()
    for path in watch_dir.iterdir():
        if not path.name.lower().endswith(".pdf"):
            continue
        name_clean = re.sub(r"[\s\-/]", "", path.stem).lower()
        if nro_clean in name_clean:
            return path
    return None


def find_pending_pdfs(config: Config, processed_hashes: set[str]) -> List[Path]:
    if not config.watch_dir.exists():
        raise FileNotFoundError(f"No existe el directorio a revisar: {config.watch_dir}")

    pdfs: List[Path] = []
    for path in sorted(config.watch_dir.iterdir()):
        if path.is_dir():
            logging.info("Omitido (es un directorio): %s", path.name)
            continue
        if not path.name.lower().endswith(".pdf"):
            logging.info("Omitido (no es un archivo PDF): %s", path.name)
            continue
        if not _VALID_PREFIXES.match(path.name):
            logging.info("Omitido (no inicia con OC, Ord. o RE): %s", path.name)
            continue
        try:
            file_hash = sha256_file(path)
        except OSError as exc:
            logging.warning("Omitido (no se pudo leer, puede estar solo en la nube): %s — %s", path.name, exc)
            continue
        if file_hash in processed_hashes:
            logging.info("Omitido (ya fue procesado anteriormente): %s", path.name)
            continue
        pdfs.append(path)
    return pdfs


_MULTA_KEYWORDS = re.compile(
    r"multa|formulaci[oó]n de cargos|cargo[s]? sancionatorio|sanci[oó]n",
    re.IGNORECASE,
)


def es_multa_o_cargos(concepto: str) -> bool:
    return bool(_MULTA_KEYWORDS.search(concepto))


@dataclass
class ProcessedOficio:
    nro: str
    categoria: str
    concepto: str
    gerencia: str
    es_multa: bool


@dataclass
class ProcessingStats:
    total: int = 0
    errores: int = 0
    categorias: Counter = field(default_factory=Counter)
    areas: Counter = field(default_factory=Counter)
    oficios_multa: List[ProcessedOficio] = field(default_factory=list)

    def registrar(self, extracted: dict[str, Any]) -> None:
        self.total += 1
        cat = extracted.get("categoria") or "Sin categoría"
        area = extracted.get("gerencia_responsable") or "Sin área"
        concepto = extracted.get("concepto") or ""
        self.categorias[cat] += 1
        self.areas[area] += 1
        if es_multa_o_cargos(concepto):
            self.oficios_multa.append(ProcessedOficio(
                nro=extracted.get("numero_oficio") or "S/N",
                categoria=cat,
                concepto=concepto,
                gerencia=area,
                es_multa=True,
            ))

    def resumen(self) -> str:
        lines = [f"PDFs nuevos procesados: {self.total}"]
        if self.errores:
            lines.append(f"Errores: {self.errores}")
        lines.append("")
        lines.append("Por categoría:")
        for cat, n in sorted(self.categorias.items()):
            lines.append(f"  {cat}: {n}")
        lines.append("")
        lines.append("Por área:")
        for area, n in sorted(self.areas.items()):
            lines.append(f"  {area}: {n}")
        if self.oficios_multa:
            lines.append("")
            lines.append(f"MULTAS / FORMULACIÓN DE CARGOS: {len(self.oficios_multa)}")
            for o in self.oficios_multa:
                lines.append(f"  !! Nro {o.nro} ({o.categoria}) — {o.gerencia}")
                lines.append(f"     {o.concepto[:120]}")
        return "\n".join(lines)


def show_summary_popup(stats: ProcessingStats) -> None:
    root = tk.Tk()
    root.withdraw()
    icon = "warning" if stats.oficios_multa else "info"
    title = "Resumen de procesamiento"
    if stats.oficios_multa:
        title += f"  —  {len(stats.oficios_multa)} MULTA(S) DETECTADA(S)"
    if icon == "warning":
        messagebox.showwarning(title, stats.resumen())
    else:
        messagebox.showinfo(title, stats.resumen())
    root.destroy()


# ---------------------------------------------------------------------------
# Informe de Multa — OpenAI + Word
# ---------------------------------------------------------------------------

def is_multa(nro: str, concepto: str, corrections: List[Dict[str, Any]]) -> bool:
    """Retorna True si el oficio es una multa según keywords o corrección manual."""
    for c in reversed(corrections):
        if str(c.get("nro", "")) == nro and c.get("campo") == "es_multa":
            return bool(c.get("valor_nuevo", False))
    return bool(_MULTA_KEYWORDS.search(concepto))


def call_anthropic_informe_multa(config: Config, pdf_path: Path) -> Dict[str, Any]:
    """Llama a la API de Anthropic (Claude) para extraer datos del informe de multa."""
    pdf_b64 = base64.b64encode(pdf_path.read_bytes()).decode("ascii")

    schema_instruction = (
        "Responde SOLO con un JSON válido que cumpla exactamente este schema:\n"
        + json.dumps(INFORME_MULTA_SCHEMA, ensure_ascii=False, indent=2)
    )

    payload = {
        "model": config.informe_multa_model,
        "max_tokens": 4096,
        "system": INFORME_MULTA_PROMPT + "\n\n" + schema_instruction,
        "messages": [
            {
                "role": "user",
                "content": [
                    {
                        "type": "document",
                        "source": {
                            "type": "base64",
                            "media_type": "application/pdf",
                            "data": pdf_b64,
                        },
                    },
                    {
                        "type": "text",
                        "text": f"Analiza este PDF de multa SEC. Archivo: {pdf_path.name}",
                    },
                ],
            }
        ],
    }
    headers = {
        "Content-Type": "application/json",
        "x-api-key": config.informe_multa_api_key,
        "anthropic-version": "2023-06-01",
    }
    resp = requests.post("https://api.anthropic.com/v1/messages", headers=headers,
                         json=payload, timeout=config.request_timeout_seconds)
    if not resp.ok:
        logging.error("Anthropic informe multa error %s: %s", resp.status_code, resp.text)
        resp.raise_for_status()

    data = resp.json()
    text_content = ""
    for block in data.get("content", []):
        if block.get("type") == "text":
            text_content += block.get("text", "")

    text_content = text_content.strip()
    if text_content.startswith("```"):
        text_content = re.sub(r"^```(?:json)?\s*", "", text_content)
        text_content = re.sub(r"\s*```$", "", text_content)

    return json.loads(text_content)


def _replace_in_paragraph(para: Any, replacements: Dict[str, str]) -> None:
    """Reemplaza placeholders {{KEY}} en un párrafo de python-docx."""
    full = "".join(r.text for r in para.runs)
    if "{{" not in full:
        return
    for key, val in replacements.items():
        full = full.replace(f"{{{{{key}}}}}", val or "")
    for i, run in enumerate(para.runs):
        run.text = full if i == 0 else ""


def create_informe_template(output_path: Path) -> None:
    """Crea el archivo Word template del informe de multa con placeholders."""
    if not DOCX_AVAILABLE:
        logging.warning("python-docx no instalado; no se puede crear el template.")
        return
    doc = DocxDocument()
    # Encabezado
    h = doc.add_heading("INFORME DE ZONA POR MULTA SEC", level=1)
    h.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Sección 1
    doc.add_heading("1. Identificación General del Evento", level=2)
    tbl = doc.add_table(rows=11, cols=2)
    tbl.style = "Table Grid"
    fields1 = [
        ("Sociedad", "{{SOCIEDAD}}"), ("Regional", "{{REGIONAL}}"),
        ("Zonal", "{{ZONAL}}"), ("Evento", "{{EVENTO}}"),
        ("Sector / Activo involucrado", "{{SECTOR_ACTIVO}}"),
        ("Clientes afectados (si aplica)", "{{CLIENTES_AFECTADOS}}"),
        ("Fecha del Evento", "{{FECHA_EVENTO}}"),
        ("Fecha de notificación de la multa SEC", "{{FECHA_NOTIFICACION}}"),
        ("Monto de la multa (UTM)", "{{MONTO_UTM}}"),
        ("N° Resolución Exenta SEC", "{{NUMERO_RESOLUCION}}"),
        ("Etapa actual de la multa", "{{ETAPA_MULTA}}"),
    ]
    for i, (label, ph) in enumerate(fields1):
        tbl.cell(i, 0).text = label
        tbl.cell(i, 1).text = ph

    # Sección 2
    doc.add_heading("2. Causa de Infracción Normativa", level=2)
    doc.add_paragraph("Descripción objetiva del motivo de la multa (según SEC):")
    doc.add_paragraph("{{DESCRIPCION_MOTIVO_SEC}}")
    doc.add_paragraph("Descripción Técnica del Evento (según Zona):")
    doc.add_paragraph("{{DESCRIPCION_TECNICA_ZONA}}")
    doc.add_paragraph("Declaración de la causa raíz interna (si se identifica):")
    doc.add_paragraph("{{CAUSA_RAIZ}}")

    # Sección 3
    doc.add_heading("3. Cronología de la sanción", level=2)
    doc.add_paragraph("{{CRONOLOGIA}}")

    # Sección 4
    doc.add_heading("4. Antecedentes Técnicos del Caso", level=2)
    doc.add_paragraph("¿Qué pasó? ¿Cuál es la problemática que se sanciona?")
    doc.add_paragraph("{{QUE_PASO}}")
    doc.add_paragraph("¿Qué se hizo y qué se dejó de hacer para que la multa fuera cursada?")
    doc.add_paragraph("{{QUE_SE_HIZO}}")
    doc.add_paragraph("Plan de mejora (si existió) y avance al momento de la multa:")
    doc.add_paragraph("{{PLAN_MEJORA}}")

    # Sección 5
    doc.add_heading("5. Área Responsable de la aplicación de la multa", level=2)
    doc.add_paragraph("{{AREA_RESPONSABLE_MULTA}}")

    # Sección 6
    doc.add_heading("6. Área Responsable de implementar el plan de mejora", level=2)
    doc.add_paragraph("{{AREA_RESPONSABLE_MEJORA}}")

    # Sección 7
    doc.add_heading("7. Firmas", level=2)
    tbl2 = doc.add_table(rows=2, cols=2)
    tbl2.cell(0, 0).text = "______________________________"
    tbl2.cell(0, 1).text = "______________________________"
    tbl2.cell(1, 0).text = "(nombre)\nDirector Regional"
    tbl2.cell(1, 1).text = "(nombre)\nDirector (Of. Central)"

    doc.add_paragraph(f"\nGenerado: {{{{FECHA_GENERACION}}}}")
    output_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(output_path))
    logging.info("Template de informe de multa creado en %s", output_path)


def fill_informe_multa(config: Config, informe_data: Dict[str, Any],
                       nro: str, output_path: Path) -> bool:
    """Rellena el template Word con los datos extraídos y lo guarda."""
    if not DOCX_AVAILABLE:
        logging.error("python-docx no instalado; no se puede generar el informe.")
        return False

    template_path = get_base_dir() / "informe_multa_template.docx"
    if not template_path.exists():
        create_informe_template(template_path)

    def v(key: str) -> str:
        return str(informe_data.get(key) or "")

    replacements = {
        "SOCIEDAD": v("sociedad"),
        "REGIONAL": v("regional"),
        "ZONAL": v("zonal"),
        "EVENTO": v("evento"),
        "SECTOR_ACTIVO": v("sector_activo"),
        "CLIENTES_AFECTADOS": v("clientes_afectados"),
        "FECHA_EVENTO": v("fecha_evento"),
        "FECHA_NOTIFICACION": v("fecha_notificacion_multa"),
        "MONTO_UTM": v("monto_utm"),
        "NUMERO_RESOLUCION": v("numero_resolucion"),
        "ETAPA_MULTA": v("etapa_multa"),
        "DESCRIPCION_MOTIVO_SEC": v("descripcion_motivo_sec"),
        "DESCRIPCION_TECNICA_ZONA": v("descripcion_tecnica_zona"),
        "CAUSA_RAIZ": v("causa_raiz"),
        "CRONOLOGIA": v("cronologia"),
        "QUE_PASO": v("que_paso"),
        "QUE_SE_HIZO": v("que_se_hizo"),
        "PLAN_MEJORA": v("plan_mejora"),
        "AREA_RESPONSABLE_MULTA": v("area_responsable_multa"),
        "AREA_RESPONSABLE_MEJORA": v("area_responsable_mejora"),
        "FECHA_GENERACION": date.today().strftime("%d-%m-%Y"),
    }

    doc = DocxDocument(str(template_path))
    for para in doc.paragraphs:
        _replace_in_paragraph(para, replacements)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    _replace_in_paragraph(para, replacements)

    output_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(output_path))
    logging.info("Informe de multa generado: %s", output_path)
    return True


def ask_and_generate_informe(config: Config, pdf_path: Path,
                              extracted: Dict[str, Any]) -> None:
    """Pregunta al usuario si desea generar el informe de multa y lo genera."""
    nro = extracted.get("numero_oficio") or "S/N"
    answer = messagebox.askyesno(
        "Informe de Multa",
        f"El oficio Nro {nro} parece ser una multa o formulación de cargos.\n\n"
        "¿Desea generar automáticamente el Informe de Zona por Multa SEC?"
    )
    if not answer:
        return
    if not DOCX_AVAILABLE:
        messagebox.showerror(
            "Dependencia faltante",
            "Instale python-docx para generar informes Word:\n  pip install python-docx"
        )
        return
    output_path = config.informe_output_dir / f"Informe_Multa_Nro{nro}.docx"
    try:
        informe_data = call_anthropic_informe_multa(config, pdf_path)
        ok = fill_informe_multa(config, informe_data, nro, output_path)
        if ok:
            messagebox.showinfo(
                "Informe generado",
                f"Informe de multa guardado en:\n{output_path}"
            )
    except Exception as exc:
        logging.exception("Error generando informe de multa: %s", exc)
        messagebox.showerror("Error", f"No se pudo generar el informe:\n{exc}")


# ---------------------------------------------------------------------------
# Microsoft Planner integration
# ---------------------------------------------------------------------------

def get_planner_token(planner: PlannerConfig) -> Optional[str]:
    authority = f"https://login.microsoftonline.com/{planner.tenant_id}"
    app = msal.ConfidentialClientApplication(
        planner.client_id,
        authority=authority,
        client_credential=planner.client_secret,
    )
    result = app.acquire_token_for_client(scopes=["https://graph.microsoft.com/.default"])
    if "access_token" in result:
        return result["access_token"]
    logging.error("No se pudo obtener token de Planner: %s", result.get("error_description", result))
    return None


def create_planner_task(token: str, planner: PlannerConfig, title: str, due_date: date) -> bool:
    headers = {
        "Authorization": f"Bearer {token}",
        "Content-Type": "application/json",
    }
    payload = {
        "planId": planner.plan_id,
        "bucketId": planner.bucket_id,
        "title": title,
        "dueDateTime": f"{due_date.isoformat()}T23:59:59Z",
    }
    resp = requests.post(
        "https://graph.microsoft.com/v1.0/planner/tasks",
        headers=headers,
        json=payload,
        timeout=30,
    )
    if resp.ok:
        logging.info("Tarea creada en Planner: %s", title)
        return True
    logging.error("Error creando tarea en Planner: %s — %s", resp.status_code, resp.text)
    return False


def sync_to_planner(config: Config, extracted: dict[str, Any], due_date: date) -> None:
    if not config.planner.enabled:
        return
    if due_date <= date.today():
        logging.info("Plazo ya vencido, no se crea tarea en Planner: %s", extracted.get("numero_oficio"))
        return
    nro = extracted.get("numero_oficio") or "S/N"
    cat = extracted.get("categoria") or ""
    concepto = extracted.get("concepto") or ""
    gerencia = extracted.get("gerencia_responsable") or ""
    title = f"[{cat}] Nro {nro} — {gerencia} — {concepto[:80]}"
    token = get_planner_token(config.planner)
    if token:
        create_planner_task(token, config.planner, title, due_date)


# ---------------------------------------------------------------------------
# Outlook attachment downloader
# ---------------------------------------------------------------------------


def find_outlook_folder_id(token: str, user_email: str, folder_name: str) -> Optional[str]:
    """Busca el ID de una carpeta de correo por nombre."""
    headers = {"Authorization": f"Bearer {token}"}
    url = f"https://graph.microsoft.com/v1.0/users/{user_email}/mailFolders"
    params = {"$filter": f"displayName eq '{folder_name}'", "$top": "50"}
    resp = requests.get(url, headers=headers, params=params, timeout=30)
    if not resp.ok:
        logging.error("Error listando carpetas de correo: %s — %s", resp.status_code, resp.text)
        return None
    folders = resp.json().get("value", [])
    for f in folders:
        if f.get("displayName") == folder_name:
            return f["id"]
    logging.warning("No se encontró la carpeta '%s' en el buzón de %s", folder_name, user_email)
    return None


def download_outlook_attachments(config: Config, state: dict[str, Any]) -> int:
    """Descarga adjuntos PDF de la carpeta Outlook configurada. Retorna cantidad de archivos nuevos."""
    if not config.outlook.enabled:
        return 0

    if not config.outlook.user_email:
        logging.warning("Outlook habilitado pero falta user_email en config.")
        return 0

    token = get_planner_token(config.planner)
    if not token:
        logging.error("No se pudo obtener token para Outlook.")
        return 0

    folder_id = find_outlook_folder_id(token, config.outlook.user_email, config.outlook.folder_name)
    if not folder_id:
        return 0

    processed_ids: set[str] = set(state.get("outlook_processed_ids", []))
    headers = {"Authorization": f"Bearer {token}"}
    user = config.outlook.user_email
    saved = 0

    page_size = 50
    skip = 0
    while True:
        url = f"https://graph.microsoft.com/v1.0/users/{user}/mailFolders/{folder_id}/messages"
        params = {
            "$top": str(page_size),
            "$skip": str(skip),
            "$select": "id,subject,hasAttachments,receivedDateTime",
            "$orderby": "receivedDateTime desc",
        }
        resp = requests.get(url, headers=headers, params=params, timeout=30)
        if not resp.ok:
            logging.error("Error listando mensajes: %s — %s", resp.status_code, resp.text)
            break

        messages = resp.json().get("value", [])
        if not messages:
            break

        all_already_processed = True
        for msg in messages:
            msg_id = msg["id"]
            if msg_id in processed_ids:
                continue
            all_already_processed = False

            if not msg.get("hasAttachments"):
                processed_ids.add(msg_id)
                continue

            att_url = f"https://graph.microsoft.com/v1.0/users/{user}/messages/{msg_id}/attachments"
            att_resp = requests.get(att_url, headers=headers, timeout=30)
            if not att_resp.ok:
                logging.error("Error obteniendo adjuntos del mensaje %s: %s", msg_id, att_resp.status_code)
                continue

            for att in att_resp.json().get("value", []):
                if att.get("@odata.type") != "#microsoft.graph.fileAttachment":
                    continue
                name = att.get("name", "")
                if not name.lower().endswith(".pdf"):
                    continue
                content_bytes = att.get("contentBytes")
                if not content_bytes:
                    continue

                dest = config.watch_dir / name
                if dest.exists():
                    logging.info("Adjunto ya existe en disco, omitido: %s", name)
                else:
                    dest.write_bytes(base64.b64decode(content_bytes))
                    logging.info("Adjunto descargado: %s (mensaje: %s)", name, msg.get("subject", ""))
                    saved += 1

            processed_ids.add(msg_id)

        if all_already_processed:
            break
        skip += page_size

    state["outlook_processed_ids"] = sorted(processed_ids)
    save_state(config.processed_state_path, state)

    if saved:
        logging.info("Se descargaron %d adjunto(s) nuevos desde Outlook.", saved)
    else:
        logging.info("No hay adjuntos nuevos en Outlook.")
    return saved


# ---------------------------------------------------------------------------
# Correcciones / aprendizaje
# ---------------------------------------------------------------------------

def load_corrections(path: Path) -> List[Dict[str, Any]]:
    if not path.exists():
        return []
    with path.open("r", encoding="utf-8") as f:
        return json.load(f)


def build_corrections_prompt(corrections: List[Dict[str, Any]]) -> str:
    if not corrections:
        return ""
    lines = [
        "\n\nCorrecciones previas del usuario (usa estos ejemplos para aprender y mejorar):"
    ]
    for c in corrections[-20:]:
        lines.append(
            f"  - Oficio Nro {c['nro']}: "
            f"campo '{c['campo']}' corregido de '{c['valor_anterior']}' a '{c['valor_nuevo']}'"
        )
        if c.get("concepto"):
            lines.append(f"    Concepto del oficio: {c['concepto'][:150]}")
    lines.append(
        "\nConsidera estas correcciones como referencia para clasificar oficios similares."
    )
    return "\n".join(lines)


def build_history_prompt(excel_path: Path, max_per_area: int = 5) -> str:
    """Lee el Excel y construye un bloque de ejemplos históricos por área.

    Selecciona hasta *max_per_area* oficios recientes de cada área para que
    el modelo aprenda los patrones de clasificación reales.
    """
    if not excel_path.exists():
        return ""

    try:
        wb = load_workbook(excel_path, read_only=True)
        ws = wb.active
    except Exception:
        return ""

    _skip = re.compile(r"solicita\s+m[aá]s\s+informaci[oó]n", re.IGNORECASE)

    # area -> lista de (nro, concepto_corto)
    by_area: Dict[str, List[Tuple[str, str]]] = {}
    for row in ws.iter_rows(min_row=2, values_only=True):
        nro = str(row[0] or "").strip()
        concepto = str(row[3] or "").strip()
        area = str(row[5] or "").strip()
        if not nro or not area or not concepto:
            continue
        if _skip.search(concepto):
            continue  # excluir oficios de solicitud de más información
        by_area.setdefault(area, []).append((nro, concepto[:120]))
    wb.close()

    if not by_area:
        return ""

    lines = [
        "\n\nEjemplos históricos de oficios clasificados correctamente por área "
        "(usa estos ejemplos para aprender los patrones de clasificación):"
    ]
    for area in sorted(by_area):
        examples = by_area[area][-max_per_area:]  # los más recientes
        lines.append(f"\n  Área «{area}»:")
        for nro, concepto in examples:
            lines.append(f"    - Nro {nro}: {concepto}")

    lines.append(
        "\nClasifica el nuevo oficio en el área que mejor se ajuste "
        "a los patrones observados arriba. Las correcciones del usuario "
        "(si las hay) tienen prioridad sobre los ejemplos históricos."
    )
    return "\n".join(lines)


# ---------------------------------------------------------------------------
# Sistema de aprendizaje: historial, reglas auto-generadas y métricas
# ---------------------------------------------------------------------------

MAX_FEWSHOT = 10
REGLAS_CADA_N = 20


def load_historial(path: Path) -> List[Dict[str, Any]]:
    """Carga el historial de clasificaciones (propuestas, correcciones)."""
    if not path.exists():
        return []
    try:
        with path.open("r", encoding="utf-8") as f:
            return json.load(f)
    except Exception as exc:
        logging.warning("Error leyendo historial %s: %s", path, exc)
        return []


def save_historial(path: Path, historial: List[Dict[str, Any]]) -> None:
    ensure_parent(path)
    with path.open("w", encoding="utf-8") as f:
        json.dump(historial, f, ensure_ascii=False, indent=2)


def add_to_historial(path: Path, extracted: Dict[str, Any],
                     area_propuesta: str, area_final: str,
                     pdf_name: str = "") -> Dict[str, Any]:
    """Agrega una entrada al historial. fue_corregido=True si area cambió."""
    historial = load_historial(path)
    entry = {
        "archivo": pdf_name,
        "numero_oficio": extracted.get("numero_oficio", ""),
        "categoria": extracted.get("categoria", ""),
        "concepto": extracted.get("concepto", ""),
        "remitente": extracted.get("remitente", ""),
        "keywords": extracted.get("keywords", []) or [],
        "confianza": extracted.get("confianza"),
        "area_propuesta": area_propuesta,
        "area_final": area_final,
        "fue_corregido": (area_propuesta or "") != (area_final or ""),
        "fecha_procesado": datetime.now().isoformat(),
    }
    historial.append(entry)
    save_historial(path, historial)
    return entry


def mark_correction_in_historial(path: Path, nro: str, new_area: str) -> None:
    """Marca como fue_corregido=True la entrada de un oficio revaluado."""
    historial = load_historial(path)
    changed = False
    for entry in historial:
        if str(entry.get("numero_oficio", "")).strip() == str(nro).strip():
            if entry.get("area_final") != new_area:
                entry["area_final"] = new_area
                entry["fue_corregido"] = True
                entry["fecha_revaluacion"] = datetime.now().isoformat()
                changed = True
    if changed:
        save_historial(path, historial)


def build_fewshot_from_historial(historial: List[Dict[str, Any]],
                                 max_examples: int = MAX_FEWSHOT) -> str:
    """Few-shot dinámico: prioriza correcciones (60/40) y los más recientes."""
    if not historial:
        return ""

    _skip = re.compile(r"solicita\s+m[aá]s\s+informaci[oó]n", re.IGNORECASE)
    filtered = [h for h in historial if not _skip.search(str(h.get("concepto", "")))]
    if not filtered:
        return ""

    correcciones = [h for h in filtered if h.get("fue_corregido")]
    confirmaciones = [h for h in filtered if not h.get("fue_corregido")]
    n_corr = min(len(correcciones), max(1, int(max_examples * 0.6)))
    n_conf = min(len(confirmaciones), max_examples - n_corr)
    selected = correcciones[-n_corr:] + confirmaciones[-n_conf:]
    if not selected:
        return ""

    lines = ["", "Ejemplos de clasificaciones anteriores (con prioridad a correcciones del usuario):"]
    for i, ex in enumerate(selected, 1):
        tag = "CORRECCIÓN" if ex.get("fue_corregido") else "Confirmado"
        lines.append(f"  [{tag}] Oficio Nro {ex.get('numero_oficio', '?')}")
        concepto = str(ex.get("concepto", ""))[:160]
        if concepto:
            lines.append(f"    Concepto: {concepto}")
        kws = ex.get("keywords") or []
        if kws:
            lines.append(f"    Keywords: {', '.join(kws[:6])}")
        if ex.get("fue_corregido"):
            lines.append(
                f"    Área propuesta (INCORRECTA): {ex.get('area_propuesta')} "
                f"→ Área correcta: {ex.get('area_final')}"
            )
        else:
            lines.append(f"    Área: {ex.get('area_final')}")
    lines.append(
        "\nUsa estos ejemplos como señal de alta prioridad, especialmente las correcciones."
    )
    return "\n".join(lines)


def load_reglas(path: Path) -> Optional[str]:
    """Carga el texto de reglas aprendidas (si existe)."""
    if not path.exists():
        return None
    try:
        with path.open("r", encoding="utf-8") as f:
            data = json.load(f)
        return data.get("reglas_texto")
    except Exception:
        return None


def save_reglas(path: Path, reglas_texto: str, stats: Dict[str, Any]) -> None:
    ensure_parent(path)
    data = {
        "generado": datetime.now().isoformat(),
        "basado_en_n_oficios": stats.get("total", 0),
        "accuracy_global": stats.get("accuracy", 0),
        "reglas_texto": reglas_texto,
    }
    with path.open("w", encoding="utf-8") as f:
        json.dump(data, f, ensure_ascii=False, indent=2)


def build_reglas_prompt(reglas_texto: Optional[str]) -> str:
    """Inserta las reglas auto-generadas en el prompt principal."""
    if not reglas_texto:
        return ""
    return (
        "\n\nReglas de clasificación APRENDIDAS (generadas automáticamente a partir del historial "
        "de decisiones del usuario — tienen prioridad sobre las reglas base cuando haya conflicto):\n"
        + reglas_texto.strip()
    )


def generate_reglas(config: Config, historial: List[Dict[str, Any]]) -> Optional[str]:
    """Usa Claude (Anthropic) para destilar reglas a partir del historial."""
    if len(historial) < 5:
        logging.info("Se requieren al menos 5 oficios para generar reglas (hay %s).",
                     len(historial))
        return None
    if not config.informe_multa_api_key or config.informe_multa_api_key.startswith("REEMPLAZAR"):
        logging.warning("No hay API key de Anthropic configurada; no se generan reglas.")
        return None

    resumen = []
    for h in historial[-200:]:  # limitar para no exceder tokens
        item = {
            "concepto": (h.get("concepto") or "")[:220],
            "keywords": h.get("keywords") or [],
            "remitente": h.get("remitente") or "",
            "area_final": h.get("area_final", ""),
            "fue_corregido": h.get("fue_corregido", False),
        }
        if h.get("fue_corregido"):
            item["area_propuesta_incorrecta"] = h.get("area_propuesta", "")
        resumen.append(item)

    prompt = (
        "Eres un analista regulatorio de CGE. Analiza el siguiente historial de "
        "clasificaciones de oficios SEC y genera REGLAS DE CLASIFICACIÓN concisas y "
        f"accionables por cada una de estas áreas: {', '.join(AREAS_VALIDAS)}.\n\n"
        "Presta ESPECIAL atención a las CORRECCIONES (donde el modelo se equivocó). "
        "Incluye reglas negativas (\"NO clasificar como X si...\") cuando sean útiles. "
        "Responde SOLO con texto plano (sin markdown ni JSON) en el formato:\n\n"
        "ÁREA: <nombre>\n- Regla 1\n- Regla 2\n\n"
        "Incluye al final una sección REGLAS GENERALES si hay patrones transversales.\n\n"
        "HISTORIAL:\n" + json.dumps(resumen, ensure_ascii=False, indent=2)
    )

    payload = {
        "model": config.informe_multa_model,
        "max_tokens": 2000,
        "messages": [{"role": "user", "content": prompt}],
    }
    headers = {
        "Content-Type": "application/json",
        "x-api-key": config.informe_multa_api_key,
        "anthropic-version": "2023-06-01",
    }
    try:
        resp = requests.post("https://api.anthropic.com/v1/messages",
                             headers=headers, json=payload,
                             timeout=config.request_timeout_seconds)
        if not resp.ok:
            logging.error("Anthropic generate_reglas error %s: %s",
                          resp.status_code, resp.text)
            return None
        data = resp.json()
        text = "".join(b.get("text", "") for b in data.get("content", [])
                       if b.get("type") == "text").strip()
        if not text:
            return None
        stats = compute_learning_stats(historial)
        save_reglas(config.reglas_path, text, stats)
        logging.info("Reglas de clasificación regeneradas (basadas en %s oficios).",
                     stats["total"])
        return text
    except Exception as exc:
        logging.exception("Error generando reglas: %s", exc)
        return None


def compute_learning_stats(historial: List[Dict[str, Any]]) -> Dict[str, Any]:
    """Calcula métricas de accuracy del agente a partir del historial."""
    if not historial:
        return {"total": 0, "correctos": 0, "accuracy": 0.0,
                "por_area": {}, "errores_frecuentes": []}

    total = len(historial)
    correctos = sum(1 for h in historial if not h.get("fue_corregido"))
    accuracy = correctos / total if total else 0.0

    por_area: Dict[str, Dict[str, Any]] = {}
    for area in AREAS_VALIDAS:
        entries = [h for h in historial if h.get("area_final") == area]
        if entries:
            ok = sum(1 for h in entries if not h.get("fue_corregido"))
            por_area[area] = {
                "total": len(entries),
                "correctos": ok,
                "accuracy": ok / len(entries),
            }

    errores = [
        f"{h.get('area_propuesta')} → {h.get('area_final')}"
        for h in historial if h.get("fue_corregido")
    ]
    errores_frecuentes = Counter(errores).most_common(5)

    return {
        "total": total,
        "correctos": correctos,
        "accuracy": accuracy,
        "por_area": por_area,
        "errores_frecuentes": errores_frecuentes,
    }


def save_corrections(path: Path, corrections: List[Dict[str, Any]]) -> None:
    ensure_parent(path)
    with path.open("w", encoding="utf-8") as f:
        json.dump(corrections, f, ensure_ascii=False, indent=2)


def update_excel_row(excel_path: Path, db_path: Path, nro: str, categoria: str,
                     updates: Dict[str, Any], gerentes: Dict[str, Gerente]) -> None:
    """Actualiza una fila existente en Excel y SQLite según Nro + Categoría."""
    db_update_oficio(db_path, nro, categoria, updates, gerentes)
    wb = load_workbook(excel_path)
    ws = wb.active
    for row in ws.iter_rows(min_row=2):
        cell_nro = str(row[0].value or "").strip()
        cell_cat = str(row[1].value or "").strip()
        if cell_nro == nro and cell_cat == categoria:
            if "gerencia_responsable" in updates:
                new_area = updates["gerencia_responsable"]
                row[5].value = new_area
                gerente = gerentes.get(new_area, Gerente(nombre=""))
                row[6].value = gerente.nombre
            if "plazo_respuesta" in updates:
                p = parse_date_yyyy_mm_dd(updates["plazo_respuesta"])
                if p:
                    row[8].value = p
                    row[8].number_format = "DD-MM-YYYY"
            if "es_multa" in updates:
                row[9].value = "Sí" if updates["es_multa"] else ""
            wb.save(excel_path)
            return
    wb.close()


# ---------------------------------------------------------------------------
# V5 Design Tokens
# ---------------------------------------------------------------------------

PAL_LIGHT = {
    "bg": "#f7f8fa", "panel": "#ffffff", "soft": "#eff2f6", "softer": "#f4f6f9",
    "border": "#e2e6ec", "borderStrong": "#d4dae3",
    "text": "#1c2633", "subtext": "#6b7684", "dim": "#9aa3b0",
    "accent": "#0B3D6B", "blue": "#1E6FB8", "blueSoft": "#e8f0fa",
    "success": "#1a7f5a", "successSoft": "#e4f3ec",
    "warn": "#c47a00", "warnSoft": "#fdf2e0",
    "danger": "#b42d2d", "dangerSoft": "#fbe7e7",
    "lilac": "#8a4fb5", "lilacSoft": "#f2ebf8",
    "teal": "#0e8a82", "tealSoft": "#dff2f0",
    "neutral": "#556270", "neutralSoft": "#eaedf1",
}

PAL_DARK = {
    "bg": "#0f1722", "panel": "#17212f", "soft": "#1c2838", "softer": "#1a2433",
    "border": "#253244", "borderStrong": "#2d3d52",
    "text": "#e4eaf2", "subtext": "#9aa8bc", "dim": "#6b7a8f",
    "accent": "#3d8bd9", "blue": "#5aa9e6", "blueSoft": "#1c3550",
    "success": "#4db87d", "successSoft": "#1a3a2a",
    "warn": "#e0a54a", "warnSoft": "#3d2e15",
    "danger": "#e06464", "dangerSoft": "#3d1e1e",
    "lilac": "#a88be8", "lilacSoft": "#2d2540",
    "teal": "#4ec9b0", "tealSoft": "#1a3530",
    "neutral": "#8798b0", "neutralSoft": "#1f2a3a",
}

AREA_COLOR_MAP = {
    "Conexiones": "blue", "PMGD": "warn",
    "Servicio al Cliente": "success", "Pérdidas": "lilac",
    "Sin área": "neutral", "Cobranza": "danger", "Lectura": "teal",
}

FONT_UI = "Segoe UI"
FONT_MONO = "Consolas"


# ---------------------------------------------------------------------------
# Data loading for the GUI
# ---------------------------------------------------------------------------

def load_oficios_for_display(config: Config) -> List[Dict[str, Any]]:
    """Reads oficios from SQLite and returns a list of dicts for the Bandeja UI."""
    if not config.db_path.exists():
        return []
    result: List[Dict[str, Any]] = []
    corrections = load_corrections(config.corrections_path)
    today = date.today()
    con = sqlite3.connect(config.db_path)
    try:
        cur = con.execute(
            "SELECT nro, categoria, fecha_oficio, concepto, gerencia, "
            "plazo_respuesta FROM oficios ORDER BY id DESC"
        )
        for nro, categoria, fecha_iso, concepto, area, plazo_iso in cur:
            nro = (nro or "").strip()
            if not nro:
                continue
            concepto = concepto or ""
            area = area or "Sin área"

            if fecha_iso:
                try:
                    fd = date.fromisoformat(fecha_iso)
                    fecha_str = fd.strftime("%d-%m-%Y")
                except ValueError:
                    fecha_str = ""
            else:
                fecha_str = ""

            if plazo_iso:
                try:
                    pd = date.fromisoformat(plazo_iso)
                    plazo_str = pd.strftime("%d-%m-%Y")
                    dias_rest = (pd - today).days
                except ValueError:
                    plazo_str = ""
                    dias_rest = None
            else:
                plazo_str = ""
                dias_rest = None

            result.append({
                "nro": nro,
                "tipo": categoria or "",
                "area": area,
                "plazo": plazo_str,
                "diasRest": dias_rest,
                "asunto": concepto,
                "multa": is_multa(nro, concepto, corrections),
                "conf": None,
                "fecha": fecha_str,
            })
    except Exception as exc:
        logging.warning("Error leyendo SQLite para GUI: %s", exc)
    finally:
        con.close()
    return result


def get_bandeja_kpis(config: Config) -> Dict[str, Any]:
    """Compute KPI values for the Bandeja header cards."""
    oficios = load_oficios_for_display(config)
    total = len(oficios)
    multas = sum(1 for o in oficios if o["multa"])
    criticos = sum(1 for o in oficios
                   if o["diasRest"] is not None and 0 <= o["diasRest"] <= 5)
    hist = load_historial(config.historial_path)
    ls = compute_learning_stats(hist)
    acc_str = f"{ls['accuracy']*100:.0f}%" if ls["total"] else "—"
    return {
        "total": total,
        "accuracy": acc_str,
        "multas": multas,
        "criticos": criticos,
        "oficios": oficios,
    }


# ---------------------------------------------------------------------------
# Processing helpers (thread-safe callbacks)
# ---------------------------------------------------------------------------

def get_upcoming_deadlines(db_path: Path, days: int = 5) -> List[Dict[str, Any]]:
    if not db_path.exists():
        return []
    result: List[Dict[str, Any]] = []
    today = date.today()
    limit = today + timedelta(days=days)
    con = sqlite3.connect(db_path)
    try:
        cur = con.execute(
            "SELECT nro, categoria, gerencia, plazo_respuesta FROM oficios "
            "WHERE plazo_respuesta IS NOT NULL AND plazo_respuesta >= ? AND plazo_respuesta <= ? "
            "ORDER BY plazo_respuesta",
            (today.isoformat(), limit.isoformat()),
        )
        for nro, cat, area, plazo_iso in cur:
            pd = date.fromisoformat(plazo_iso)
            result.append({
                "nro": nro or "",
                "categoria": cat or "",
                "area": area or "",
                "plazo_str": pd.strftime("%d-%m-%Y"),
                "dias": (pd - today).days,
            })
    except Exception:
        pass
    finally:
        con.close()
    return result


def process_directory(
    config: Config, state: dict[str, Any]
) -> Tuple[ProcessingStats, List[Tuple[Path, Dict[str, Any]]]]:
    """Procesa los PDFs pendientes. Retorna (stats, lista_de_multas)."""
    empty_stats = ProcessingStats()
    ensure_excel_exists(config.excel_path)
    init_db(config.db_path)
    migrate_excel_to_db(config.excel_path, config.db_path)
    remove_duplicate_files(config.watch_dir, config.scan_extensions)
    processed_hashes: set[str] = set(state.get("processed_hashes", []))
    pending = find_pending_pdfs(config, processed_hashes)

    if not pending:
        logging.info("No hay PDFs nuevos para procesar en %s", config.watch_dir)
        return empty_stats, []

    logging.info("Se encontraron %s PDF(s) nuevos para procesar.", len(pending))

    corrections = load_corrections(config.corrections_path)
    corrections_prompt = build_corrections_prompt(corrections)
    history_prompt = build_history_prompt(config.excel_path)
    historial = load_historial(config.historial_path)
    fewshot_prompt = build_fewshot_from_historial(historial)
    reglas_prompt = build_reglas_prompt(load_reglas(config.reglas_path))
    learning_prompt = history_prompt + fewshot_prompt + corrections_prompt + reglas_prompt

    stats = ProcessingStats()
    multa_pdfs: List[Tuple[Path, Dict[str, Any]]] = []
    changed = False
    procesados_esta_sesion = 0
    for pdf_path in pending:
        logging.info("Procesando %s", pdf_path.name)
        file_hash = sha256_file(pdf_path)
        try:
            extracted = call_openai_extract(config, pdf_path, learning_prompt)

            related_nro = extracted.get("oficio_relacionado")
            if related_nro:
                related_path = find_related_pdf(config.watch_dir, related_nro)
                if related_path and related_path != pdf_path:
                    logging.info("Oficio relacionado encontrado: %s → re-extrayendo.",
                                 related_path.name)
                    extracted = call_openai_extract(
                        config, pdf_path, learning_prompt, related_pdf_path=related_path
                    )
                elif related_nro:
                    logging.info("Oficio relacionado Nro %s no encontrado.", related_nro)

            due_date = compute_due_date(extracted)
            if due_date and not extracted.get("plazo_respuesta"):
                logging.info("Plazo relativo calculado: %s → %s", pdf_path.name, due_date)
            row = map_row(extracted, config.gerentes)
            append_to_excel(config.excel_path, row)
            db_insert_oficio(config.db_path, row)
            if due_date:
                sync_to_planner(config, extracted, due_date)
            stats.registrar(extracted)
            nro = extracted.get("numero_oficio") or ""
            concepto = extracted.get("concepto") or ""
            area_propuesta = extracted.get("gerencia_responsable") or ""
            add_to_historial(
                config.historial_path, extracted,
                area_propuesta=area_propuesta,
                area_final=area_propuesta,
                pdf_name=pdf_path.name,
            )
            if is_multa(nro, concepto, corrections):
                multa_pdfs.append((pdf_path, extracted))
            processed_hashes.add(file_hash)
            changed = True
            procesados_esta_sesion += 1
            logging.info("PDF procesado: %s (confianza=%s)", pdf_path.name,
                         extracted.get("confianza"))
        except Exception as exc:
            stats.errores += 1
            logging.exception("Error procesando %s: %s", pdf_path.name, exc)

    if changed:
        state["processed_hashes"] = sorted(processed_hashes)
        save_state(config.processed_state_path, state)

    try:
        historial_actual = load_historial(config.historial_path)
        total_hist = len(historial_actual)
        reglas_existen = config.reglas_path.exists()
        if procesados_esta_sesion > 0 and total_hist >= REGLAS_CADA_N and (
            not reglas_existen
            or (total_hist % REGLAS_CADA_N) < procesados_esta_sesion
        ):
            logging.info("Regenerando reglas (total=%s)", total_hist)
            generate_reglas(config, historial_actual)
    except Exception as exc:
        logging.warning("No se pudieron regenerar reglas: %s", exc)

    return stats, multa_pdfs


def parse_run_time(run_time: str) -> tuple[int, int]:
    hour_str, minute_str = run_time.split(":")
    return int(hour_str), int(minute_str)


def service_loop(config: Config) -> None:
    tz = ZoneInfo(config.timezone)
    logging.info("Servicio iniciado. tz=%s | hora=%s", config.timezone, config.run_time)
    while True:
        state = load_state(config.processed_state_path)
        now = datetime.now(tz)
        hour, minute = parse_run_time(config.run_time)
        today = now.date().isoformat()
        if (now.hour > hour or (now.hour == hour and now.minute >= minute)) \
                and state.get("last_run_date") != today:
            logging.info("Ejecución diaria iniciada.")
            stats, multa_pdfs = process_directory(config, state)
            state["last_run_date"] = today
            save_state(config.processed_state_path, state)
            logging.info("Ejecución diaria finalizada.")
        time.sleep(30)


def run_once(config: Config) -> None:
    state = load_state(config.processed_state_path)
    process_directory(config, state)


def reset_state(config: Config) -> None:
    state = {"processed_hashes": [], "last_run_date": None}
    save_state(config.processed_state_path, state)
    logging.info("Estado reseteado.")


# ---------------------------------------------------------------------------
# V5 CustomTkinter UI
# ---------------------------------------------------------------------------

def _hex(color: str) -> str:
    return color


def _font(size: int, weight: str = "normal") -> Tuple[str, int, str]:
    if weight == "bold":
        return (FONT_UI, size, "bold")
    if weight == "semibold":
        return (FONT_UI, size, "bold")
    return (FONT_UI, size)


def _mono(size: int, weight: str = "normal") -> Tuple[str, int, str]:
    if weight in ("bold", "semibold"):
        return (FONT_MONO, size, "bold")
    return (FONT_MONO, size)


def _card(parent, pal: Dict, **kw) -> "ctk.CTkFrame":
    return ctk.CTkFrame(
        parent,
        fg_color=pal["panel"],
        border_width=1,
        border_color=pal["border"],
        corner_radius=10,
        **kw,
    )


def _tframe(parent, bg: str, **kw) -> tk.Frame:
    """Lightweight transparent-equivalent frame using native tkinter."""
    return tk.Frame(parent, bg=bg, **kw)


class _OficioCard:
    """Reusable card widget. Build once, update via .update(o); pool-friendly."""

    def __init__(self, app: "OficiosApp", parent) -> None:
        self.app = app
        pal = app.pal
        pbg = pal["panel"]

        self.frame = _card(parent, pal)
        self.frame.columnconfigure(0, weight=1)

        # ── Top row: nro · tipo · chips ─────────────────────
        top = tk.Frame(self.frame, bg=pbg)
        top.pack(fill="x", padx=16, pady=(14, 0))

        self.lbl_nro = tk.Label(top, fg=pal["text"], bg=pbg,
                                font=_mono(12, "bold"))
        self.lbl_nro.pack(side="left")
        self.lbl_tipo = tk.Label(top, fg=pal["subtext"], bg=pbg, font=_font(11))
        self.lbl_tipo.pack(side="left", padx=(4, 0))

        self.chip_dias = ctk.CTkLabel(
            top, text="", font=_font(10, "bold"),
            corner_radius=4, padx=8, pady=2,
        )
        self.chip_dias.pack(side="right")

        self.chip_multa = ctk.CTkLabel(
            top, text="MULTA", text_color=pal["warn"],
            fg_color=pal["warnSoft"], font=_font(10, "bold"),
            corner_radius=4, padx=8, pady=2,
        )
        self._multa_shown = False

        # ── Asunto ──────────────────────────────────────────
        self.lbl_asunto = tk.Label(
            self.frame, fg=pal["text"], bg=pbg, font=_font(13),
            anchor="w", justify="left", wraplength=480,
        )
        self.lbl_asunto.pack(fill="x", padx=16, pady=(8, 0))

        # ── Separator ───────────────────────────────────────
        tk.Frame(self.frame, bg=pal["border"], height=1).pack(
            fill="x", padx=16, pady=(10, 0))

        # ── Footer ──────────────────────────────────────────
        footer = tk.Frame(self.frame, bg=pbg)
        footer.pack(fill="x", padx=16, pady=(8, 14))

        self.dot_area = tk.Canvas(footer, width=10, height=10,
                                  highlightthickness=0, bd=0)
        self.dot_area.pack(side="left", pady=4)
        self.lbl_area = tk.Label(footer, fg=pal["text"], bg=pbg, font=_font(11))
        self.lbl_area.pack(side="left", padx=(6, 0))
        self.lbl_plazo = tk.Label(footer, fg=pal["subtext"], bg=pbg, font=_font(11))
        self.lbl_plazo.pack(side="left", padx=(4, 0))
        self._plazo_packed = True

        # Buttons via grid so MULTA toggle keeps order stable
        btn_wrap = tk.Frame(footer, bg=pbg)
        btn_wrap.pack(side="right")

        self.btn_abrir = ctk.CTkButton(
            btn_wrap, text="Abrir",
            fg_color=pal["accent"], hover_color=pal["blue"],
            text_color="#ffffff", border_width=0,
            font=_font(11), height=26, width=60, corner_radius=6,
        )
        self.btn_abrir.grid(row=0, column=0)

        self.btn_revaluar = ctk.CTkButton(
            btn_wrap, text="Revaluar",
            fg_color="transparent", hover_color=pal["soft"],
            text_color=pal["text"], border_width=1,
            border_color=pal["border"],
            font=_font(11), height=26, width=74, corner_radius=6,
        )
        self.btn_revaluar.grid(row=0, column=1, padx=(6, 0))

        self.btn_informe = ctk.CTkButton(
            btn_wrap, text="Informe",
            fg_color="transparent", hover_color=pal["soft"],
            text_color=pal["text"], border_width=1,
            border_color=pal["border"],
            font=_font(11), height=26, width=70, corner_radius=6,
        )
        self.btn_informe.grid(row=0, column=2, padx=(6, 0))
        self.btn_informe.grid_remove()
        self._informe_shown = False

    def update(self, o: Dict[str, Any]) -> None:
        app = self.app
        pal = app.pal
        dias = o["diasRest"]
        if dias is None:
            tc = pal["subtext"]; bg_chip = pal["soft"]
        elif dias <= 3:
            tc = pal["danger"]; bg_chip = pal["dangerSoft"]
        elif dias <= 5:
            tc = pal["warn"]; bg_chip = pal["warnSoft"]
        else:
            tc = pal["success"]; bg_chip = pal["successSoft"]

        area_color = pal[AREA_COLOR_MAP.get(o["area"], "blue")]

        self.lbl_nro.config(text=o["nro"])
        self.lbl_tipo.config(text=f"· {o['tipo']}")
        self.chip_dias.configure(
            text=f"{dias}d" if dias is not None else "—",
            text_color=tc, fg_color=bg_chip,
        )

        if o["multa"]:
            if not self._multa_shown:
                self.chip_multa.pack(side="right", padx=(0, 6))
                self._multa_shown = True
        elif self._multa_shown:
            self.chip_multa.pack_forget()
            self._multa_shown = False

        asunto = o["asunto"][:130] + ("…" if len(o["asunto"]) > 130 else "")
        self.lbl_asunto.config(text=asunto)

        self.dot_area.config(bg=area_color)
        self.lbl_area.config(text=o["area"])

        if o["plazo"]:
            self.lbl_plazo.config(text=f"· {o['plazo']}")
            if not self._plazo_packed:
                self.lbl_plazo.pack(side="left", padx=(4, 0))
                self._plazo_packed = True
        elif self._plazo_packed:
            self.lbl_plazo.pack_forget()
            self._plazo_packed = False

        self.btn_abrir.configure(command=lambda nro=o["nro"]: app._open_oficio(nro))
        self.btn_revaluar.configure(command=lambda oficio=o: app._go("revaluar", oficio))
        if o["multa"]:
            self.btn_informe.configure(command=lambda oficio=o: app._go("multa", oficio))
            if not self._informe_shown:
                self.btn_informe.grid()
                self._informe_shown = True
        elif self._informe_shown:
            self.btn_informe.grid_remove()
            self._informe_shown = False

    def show(self, row: int, col: int, pad_left: int) -> None:
        self.frame.grid(row=row, column=col, sticky="nsew",
                        padx=(pad_left, 0), pady=(0, 12))

    def hide(self) -> None:
        self.frame.grid_forget()


class OficiosApp(ctk.CTk):
    """Main application window — V5 Minimal design."""

    def __init__(self, config: Config) -> None:
        super().__init__()
        self.config = config
        self.pal = PAL_LIGHT
        self._dark = False
        self._screen = "inicio"
        self._selected_oficio: Optional[Dict[str, Any]] = None
        self._status_msg = "Sistema listo"
        self._last_run_info = ""
        self._cached_oficios: Optional[List[Dict[str, Any]]] = None
        self._cached_kpis: Optional[Dict[str, Any]] = None
        self._search_after_id: Optional[str] = None
        self._MAX_CARDS = 20
        self._card_pool: List[_OficioCard] = []
        self._empty_widget: Optional[tk.Frame] = None
        self._more_widget: Optional[tk.Label] = None
        self._batch_after_id: Optional[str] = None
        self._pending_render: Optional[Tuple] = None
        self._render_generation: int = 0

        ctk.set_appearance_mode("light")
        ctk.set_default_color_theme("blue")
        ctk.set_widget_scaling(1.0)
        ctk.set_window_scaling(1.0)

        self.title("Gestión de Oficios CGE")
        self.geometry("1240x820")
        self.minsize(900, 600)
        self.configure(fg_color=self.pal["bg"])

        self._build_chrome()
        self._header_frame = self._build_header()
        self._content_frame = tk.Frame(self, bg=self.pal["bg"])
        self._content_frame.pack(fill="both", expand=True)
        self._footer_frame = self._build_footer()

        self._reload_data()
        self._show_bandeja()

    # ── Window chrome (macOS-style dots) ─────────────────────────────────

    def _build_chrome(self) -> None:
        bg = self.pal["soft"]
        chrome = tk.Frame(self, bg=bg, height=32)
        chrome.pack(fill="x", side="top")
        chrome.pack_propagate(False)
        dot_frame = tk.Frame(chrome, bg=bg)
        dot_frame.pack(side="left", padx=14)
        for color in ("#ed6a5e", "#f5bf4f", "#62c554"):
            tk.Canvas(dot_frame, bg=color, width=11, height=11,
                      highlightthickness=0).pack(side="left", padx=3, pady=10)
        tk.Label(chrome, text="Gestión de Oficios CGE",
                 fg=self.pal["subtext"], bg=bg,
                 font=_font(12)).pack(expand=True)

    # ── Header ────────────────────────────────────────────────────────────

    def _build_header(self) -> tk.Frame:
        pal = self.pal
        bg = pal["panel"]
        hdr = tk.Frame(self, bg=bg, height=72)
        hdr.pack(fill="x")
        hdr.pack_propagate(False)

        inner = tk.Frame(hdr, bg=bg)
        inner.pack(fill="both", expand=True, padx=28, pady=10)

        # Logo
        logo = ctk.CTkFrame(inner, fg_color=pal["accent"],
                             width=32, height=32, corner_radius=7)
        logo.pack(side="left")
        logo.pack_propagate(False)
        tk.Label(logo, text="CGE", fg="#ffffff", bg=pal["accent"],
                 font=(FONT_UI, 12, "bold")).pack(expand=True)

        # Title block
        title_block = tk.Frame(inner, bg=bg)
        title_block.pack(side="left", padx=(10, 0))
        tk.Label(title_block, text="Gestión de Oficios",
                 fg=pal["text"], bg=bg,
                 font=_font(15, "semibold")).pack(anchor="w")
        tk.Label(title_block, text="Comercial · Servicio al Cliente",
                 fg=pal["subtext"], bg=bg,
                 font=_font(11)).pack(anchor="w")

        # Right side
        right = tk.Frame(inner, bg=bg)
        right.pack(side="right")

        # Ejecutar button
        self._btn_run = ctk.CTkButton(
            right, text="▶  Ejecutar análisis",
            fg_color=pal["accent"], hover_color=pal["blue"],
            text_color="#ffffff", font=_font(12, "semibold"),
            height=32, corner_radius=6,
            command=self._on_run,
        )
        self._btn_run.pack(side="right", padx=(10, 0))

        # Reset + re-analyze button
        ctk.CTkButton(
            right, text="⟳  Desde cero",
            fg_color="transparent", hover_color=pal["soft"],
            text_color=pal["text"], border_width=1,
            border_color=pal["border"],
            font=_font(12), height=32, corner_radius=6,
            command=self._on_reset_and_run,
        ).pack(side="right", padx=(6, 0))

        # Tab nav (segmented control)
        nav_bg = ctk.CTkFrame(right, fg_color=pal["soft"], corner_radius=8)
        nav_bg.pack(side="right")

        self._tab_btns: Dict[str, ctk.CTkButton] = {}
        for tab_id, tab_label in [("inicio", "Bandeja"), ("stats", "Estadísticas")]:
            btn = ctk.CTkButton(
                nav_bg, text=tab_label,
                fg_color="transparent", hover_color=pal["border"],
                text_color=pal["text"],
                font=_font(13), height=28, width=100, corner_radius=6,
                command=lambda tid=tab_id: self._go(tid),
            )
            btn.pack(side="left", padx=3, pady=3)
            self._tab_btns[tab_id] = btn
        self._update_tab_styles()

        # Bottom border line
        tk.Frame(self, bg=pal["border"], height=1).pack(fill="x")
        return hdr

    def _update_tab_styles(self) -> None:
        for tid, btn in self._tab_btns.items():
            active = (tid == self._screen or
                      (tid == "inicio" and self._screen in ("revaluar", "multa")))
            btn.configure(
                fg_color=self.pal["panel"] if active else "transparent",
                font=_font(13, "semibold" if active else "normal"),
            )

    # ── Footer ────────────────────────────────────────────────────────────

    def _build_footer(self) -> tk.Frame:
        pal = self.pal
        bg = pal["panel"]
        tk.Frame(self, bg=pal["border"], height=1).pack(fill="x", side="bottom")

        ftr = tk.Frame(self, bg=bg, height=34)
        ftr.pack(fill="x", side="bottom")
        ftr.pack_propagate(False)

        inner = tk.Frame(ftr, bg=bg)
        inner.pack(fill="both", expand=True, padx=28)

        tk.Canvas(inner, bg=pal["success"], width=7, height=7,
                  highlightthickness=0).pack(side="left", pady=13)

        self._status_label = tk.Label(
            inner, text=self._status_msg,
            fg=pal["subtext"], bg=bg, font=_font(11),
        )
        self._status_label.pack(side="left", padx=(8, 0))

        tk.Label(inner, text="v2.5.0",
                 fg=pal["dim"], bg=bg, font=_font(10)).pack(side="right")
        return ftr

    def _set_status(self, msg: str, dot_color: Optional[str] = None) -> None:
        self.after(0, lambda: self._status_label.config(text=msg))

    # ── Data caching ─────────────────────────────────────────────────────

    def _reload_data(self) -> None:
        """Sync load — used at startup when we must have data before render."""
        kpis = get_bandeja_kpis(self.config)
        self._cached_kpis = kpis
        self._cached_oficios = kpis["oficios"]

    def _reload_data_async(self, on_done: Optional[Callable[[], None]] = None) -> None:
        """Load Excel data in a background thread. Calls on_done on the UI
        thread when finished (or on error, with status updated)."""
        def _worker() -> None:
            try:
                kpis = get_bandeja_kpis(self.config)
                self.after(0, lambda: self._apply_reloaded(kpis, on_done))
            except Exception as exc:
                logging.exception("Error cargando datos: %s", exc)
                self.after(0, lambda e=exc: self._set_status(f"✗ Error al cargar: {e}"))
        threading.Thread(target=_worker, daemon=True).start()

    def _apply_reloaded(self, kpis: Dict[str, Any],
                        on_done: Optional[Callable[[], None]]) -> None:
        self._cached_kpis = kpis
        self._cached_oficios = kpis["oficios"]
        if on_done is not None:
            on_done()

    # ── Screen switching ──────────────────────────────────────────────────

    def _go(self, screen: str, payload: Optional[Dict] = None) -> None:
        self._screen = screen
        if screen == "revaluar":
            self._selected_oficio = payload
        elif screen == "multa":
            self._selected_oficio = payload
        # Cancel any in-flight progressive render before tearing down
        if self._batch_after_id is not None:
            self.after_cancel(self._batch_after_id)
            self._batch_after_id = None
        self._pending_render = None
        self._render_generation += 1
        self._card_pool = []
        self._empty_widget = None
        self._more_widget = None
        self._update_tab_styles()
        for w in self._content_frame.winfo_children():
            w.destroy()
        if screen == "inicio":
            self._show_bandeja()
        elif screen == "stats":
            self._show_stats_placeholder()
        elif screen == "revaluar":
            self._show_revaluar_placeholder()
        elif screen == "multa":
            self._show_multa_placeholder()

    def _show_stats_placeholder(self) -> None:
        bg = self.pal["bg"]
        f = _tframe(self._content_frame, bg)
        f.pack(fill="both", expand=True, padx=28, pady=28)
        tk.Label(f, text="Estadísticas — próxima iteración",
                 fg=self.pal["subtext"], bg=bg, font=_font(14)).pack(expand=True)

    def _show_revaluar_placeholder(self) -> None:
        bg = self.pal["bg"]
        f = _tframe(self._content_frame, bg)
        f.pack(fill="both", expand=True, padx=28, pady=28)
        nro = (self._selected_oficio or {}).get("nro", "")
        tk.Label(f, text=f"Revaluar oficio {nro} — próxima iteración",
                 fg=self.pal["subtext"], bg=bg, font=_font(14)).pack(expand=True)

    def _show_multa_placeholder(self) -> None:
        bg = self.pal["bg"]
        f = _tframe(self._content_frame, bg)
        f.pack(fill="both", expand=True, padx=28, pady=28)
        nro = (self._selected_oficio or {}).get("nro", "")
        tk.Label(f, text=f"Informe de multa {nro} — próxima iteración",
                 fg=self.pal["subtext"], bg=bg, font=_font(14)).pack(expand=True)

    # ── Ejecutar analysis ─────────────────────────────────────────────────

    def _on_run(self) -> None:
        self._btn_run.configure(state="disabled", text="⏳  Procesando…")
        self._set_status("⏳ Procesando PDFs, por favor espere…")

        def _worker() -> None:
            try:
                state = load_state(self.config.processed_state_path)
                stats, multa_pdfs = process_directory(self.config, state)
                self.after(0, lambda: self._on_run_done(stats, multa_pdfs))
            except Exception as exc:
                self.after(0, lambda e=exc: self._on_run_error(e))

        threading.Thread(target=_worker, daemon=True).start()

    def _on_run_done(self, stats: ProcessingStats,
                     multa_pdfs: List[Tuple[Path, Dict[str, Any]]]) -> None:
        self._btn_run.configure(state="normal", text="▶  Ejecutar análisis")
        msg = (f"✓ Completado: {stats.total} PDF(s) procesado(s)"
               if stats.total else "● Sin PDFs nuevos")
        self._set_status(msg)

        def _after_reload() -> None:
            if self._screen == "inicio":
                self._go("inicio")

        self._reload_data_async(on_done=_after_reload)
        for pdf_path, extracted in multa_pdfs:
            ask_and_generate_informe(self.config, pdf_path, extracted)

    def _on_run_error(self, exc: Exception) -> None:
        self._btn_run.configure(state="normal", text="▶  Ejecutar análisis")
        self._set_status(f"✗ Error: {exc}")
        messagebox.showerror("Error al procesar", str(exc))

    def _on_reset_and_run(self) -> None:
        if not messagebox.askyesno(
            "Confirmar reset",
            "Esto borrará la memoria de PDFs procesados y re-analizará todo desde cero.\n\n¿Continuar?",
        ):
            return
        reset_state(self.config)
        self._set_status("🔄 Memoria reseteada — re-analizando…")
        self._on_run()

    def _open_oficio(self, nro: str) -> None:
        pdf_path = find_related_pdf(self.config.watch_dir, nro)
        if pdf_path is None:
            self._set_status(f"✗ No se encontró PDF para oficio {nro}")
            return
        try:
            if hasattr(os, "startfile"):
                os.startfile(pdf_path)
            else:
                import subprocess
                subprocess.Popen(["xdg-open", str(pdf_path)])
        except Exception as exc:
            self._set_status(f"✗ Error abriendo PDF: {exc}")

    # ── Bandeja ───────────────────────────────────────────────────────────

    def _show_bandeja(self) -> None:
        pal = self.pal
        if self._cached_kpis is None:
            self._reload_data()
        kpis = self._cached_kpis
        oficios = self._cached_oficios

        bg = pal["bg"]

        # Outer scroll container
        scroll = ctk.CTkScrollableFrame(
            self._content_frame, fg_color=bg, corner_radius=0,
            scrollbar_button_color=pal["border"],
            scrollbar_button_hover_color=pal["borderStrong"],
        )
        scroll.pack(fill="both", expand=True)

        inner = _tframe(scroll, bg)
        inner.pack(fill="both", expand=True, padx=28, pady=(20, 28))

        # ── KPI row ───────────────────────────────────────────────────────
        kpi_row = _tframe(inner, bg)
        kpi_row.pack(fill="x", pady=(0, 14))
        for i in range(4):
            kpi_row.columnconfigure(i, weight=1, uniform="kpi")

        kpi_data = [
            (str(kpis["total"]), "Oficios totales",
             f"+{sum(1 for o in oficios if o['fecha'] == date.today().strftime('%d-%m-%Y'))} hoy",
             pal["text"], "▦"),
            (kpis["accuracy"], "Accuracy agente",
             "historial del agente", pal["success"], "✓"),
            (str(kpis["multas"]), "Multas detectadas",
             "formulaciones de cargos", pal["warn"], "◆"),
            (str(kpis["criticos"]), "Plazos críticos",
             "menos de 5 días", pal["danger"], "⏱"),
        ]
        for col, (val, label, sub, color, icon) in enumerate(kpi_data):
            self._kpi_card(kpi_row, val, label, sub, color, icon, col)

        # ── Alert bar ─────────────────────────────────────────────────────
        criticos_list = [o for o in oficios
                         if o["diasRest"] is not None and 0 <= o["diasRest"] <= 5]
        if criticos_list:
            alert_bg = pal["warnSoft"]
            alert = tk.Frame(inner, bg=alert_bg)
            alert.pack(fill="x", pady=(0, 14))
            alert_inner = tk.Frame(alert, bg=alert_bg)
            alert_inner.pack(fill="x", padx=16, pady=12)

            tk.Label(alert_inner, text="⚠", fg=pal["warn"], bg=alert_bg,
                     font=_font(16, "bold")).pack(side="left")

            txt_f = tk.Frame(alert_inner, bg=alert_bg)
            txt_f.pack(side="left", padx=(12, 0), fill="x", expand=True)
            n = len(criticos_list)
            tk.Label(txt_f,
                     text=f"{n} oficio{'s' if n > 1 else ''} con plazo en menos de 5 días",
                     fg=pal["text"], bg=alert_bg, font=_font(13, "semibold"),
                     anchor="w").pack(anchor="w")
            detail = "  ·  ".join(
                f"{o['nro']} vence en {o['diasRest']}d"
                + (" (multa)" if o["multa"] else "")
                for o in criticos_list[:4]
            )
            tk.Label(txt_f, text=detail,
                     fg=pal["subtext"], bg=alert_bg, font=_font(11),
                     anchor="w").pack(anchor="w", pady=(2, 0))

        # ── Filter tabs + search ──────────────────────────────────────────
        self._bandeja_tab = tk.StringVar(value="hoy")
        self._bandeja_q = tk.StringVar(value="")

        tab_bar_outer = _tframe(inner, bg)
        tab_bar_outer.pack(fill="x", pady=(0, 2))

        tabs_frame = _tframe(tab_bar_outer, bg)
        tabs_frame.pack(side="left")

        multas_count = sum(1 for o in oficios if o["multa"])
        vencer_count = sum(1 for o in oficios
                           if o["diasRest"] is not None and 0 <= o["diasRest"] <= 5)
        tab_defs = [
            ("hoy",     "Hoy",       len(oficios),  None),
            ("vencer",  "Por vencer", vencer_count,  "warn"),
            ("multas",  "Multas",    multas_count,   "danger"),
            ("todos",   "Histórico", len(oficios),   None),
        ]

        self._tab_filter_btns: Dict[str, Tuple] = {}
        for tid, tlabel, tcount, tone in tab_defs:
            btn = self._tab_btn(tabs_frame, tid, tlabel, tcount, tone,
                                lambda t=tid: self._set_bandeja_tab(t))
            btn.pack(side="left")
        self._set_bandeja_tab("hoy", refresh=False)

        # Search
        search_wrap = tk.Frame(tab_bar_outer, bg=pal["panel"],
                               highlightbackground=pal["border"], highlightthickness=1)
        search_wrap.pack(side="right")
        tk.Label(search_wrap, text="⌕", fg=pal["subtext"], bg=pal["panel"],
                 font=_font(12)).pack(side="left", padx=(8, 2))
        search_entry = tk.Entry(
            search_wrap, textvariable=self._bandeja_q,
            width=28, font=_font(12), bd=0, bg=pal["panel"],
            fg=pal["text"], insertbackground=pal["text"],
        )
        search_entry.pack(side="left", padx=(0, 8), pady=4)
        self._bandeja_q.trace_add("write", lambda *_: self._debounced_refresh())

        # Tab underline
        tk.Frame(inner, bg=pal["border"], height=1).pack(fill="x", pady=(0, 14))

        # ── Cards grid ────────────────────────────────────────────────────
        self._cards_container = _tframe(inner, bg)
        self._cards_container.pack(fill="both", expand=True)
        self._cards_container.columnconfigure(0, weight=1, uniform="card")
        self._cards_container.columnconfigure(1, weight=1, uniform="card")

        # Pre-build the pool of reusable cards (hidden). Since
        # _cards_container is rebuilt each _show_bandeja, the pool too.
        self._card_pool = [
            _OficioCard(self, self._cards_container)
            for _ in range(self._MAX_CARDS)
        ]

        self._empty_widget = tk.Frame(
            self._cards_container, bg=bg,
            highlightbackground=pal["border"], highlightthickness=1,
        )
        tk.Label(self._empty_widget, text="Sin oficios para este filtro.",
                 fg=pal["subtext"], bg=bg, font=_font(13)).pack(pady=40)

        self._more_widget = tk.Label(
            self._cards_container, text="",
            fg=pal["subtext"], bg=bg, font=_font(12),
        )

        self._all_oficios = oficios
        self._refresh_cards()

    def _kpi_card(self, parent, val: str, label: str, sub: str,
                  color: str, icon: str, col: int) -> None:
        pal = self.pal
        card = _card(parent, pal)
        card.grid(row=0, column=col, padx=(0 if col == 0 else 6, 0), sticky="nsew")

        panel_bg = pal["panel"]
        tk.Label(card, text=icon, fg=pal["dim"], bg=panel_bg,
                 font=_font(18), anchor="e").place(relx=1.0, rely=0.0, x=-14, y=14)

        body = tk.Frame(card, bg=panel_bg)
        body.pack(padx=16, pady=14, anchor="w")
        tk.Label(body, text=label, fg=pal["subtext"], bg=panel_bg,
                 font=_font(11), anchor="w").pack(anchor="w")
        tk.Label(body, text=val, fg=color, bg=panel_bg,
                 font=_font(28, "bold"), anchor="w").pack(anchor="w", pady=(2, 0))
        tk.Label(body, text=sub,
                 fg=color if color != pal["text"] else pal["subtext"],
                 bg=panel_bg, font=_font(11), anchor="w").pack(anchor="w", pady=(2, 0))

    def _tab_btn(self, parent, tid: str, label: str, count: int,
                 tone: Optional[str], command) -> tk.Frame:
        pal = self.pal
        bg = pal["bg"]
        frame = tk.Frame(parent, bg=bg, cursor="hand2")
        lbl = tk.Label(frame, text=f"{label}  {count}",
                       fg=pal["subtext"], bg=bg, font=_font(13),
                       padx=10, pady=8, cursor="hand2")
        lbl.pack()
        lbl.bind("<Button-1>", lambda e: command())
        underline = tk.Frame(frame, bg=bg, height=2)
        underline.pack(fill="x")
        self._tab_filter_btns[tid] = (lbl, underline)
        return frame

    def _set_bandeja_tab(self, tid: str, refresh: bool = True) -> None:
        pal = self.pal
        self._bandeja_tab.set(tid)
        for t, (lbl, underline) in self._tab_filter_btns.items():
            active = t == tid
            lbl.config(
                fg=pal["text"] if active else pal["subtext"],
                font=_font(13, "semibold" if active else "normal"),
            )
            underline.config(bg=pal["accent"] if active else pal["bg"])
        if refresh:
            self._refresh_cards()

    def _debounced_refresh(self) -> None:
        if self._search_after_id is not None:
            self.after_cancel(self._search_after_id)
        self._search_after_id = self.after(300, self._refresh_cards)

    def _refresh_cards(self) -> None:
        self._search_after_id = None
        if not self._card_pool:
            return

        # Cancel any in-flight progressive render
        if self._batch_after_id is not None:
            self.after_cancel(self._batch_after_id)
            self._batch_after_id = None
        self._render_generation += 1

        tab = self._bandeja_tab.get()
        q = self._bandeja_q.get().lower()
        oficios = self._all_oficios

        if tab == "vencer":
            oficios = [o for o in oficios
                       if o["diasRest"] is not None and 0 <= o["diasRest"] <= 5]
        elif tab == "multas":
            oficios = [o for o in oficios if o["multa"]]

        if q:
            oficios = [o for o in oficios
                       if q in (o["nro"] + o["asunto"] + o["area"]).lower()]

        total_filtered = len(oficios)
        visible = oficios[:self._MAX_CARDS]

        # Empty state + hide all extras up-front (cheap)
        if self._empty_widget is not None:
            if not visible:
                self._empty_widget.grid(row=0, column=0, columnspan=2, sticky="nsew")
            else:
                self._empty_widget.grid_forget()
        for idx in range(len(visible), len(self._card_pool)):
            self._card_pool[idx].hide()
        if not visible:
            if self._more_widget is not None:
                self._more_widget.grid_forget()
            return

        # Render first batch synchronously (above-the-fold)
        FIRST_BATCH = 6
        first_n = min(FIRST_BATCH, len(visible))
        for idx in range(first_n):
            self._render_card_at(idx, visible[idx])

        if first_n >= len(visible):
            self._finalize_render(total_filtered, len(visible))
            return

        # Defer remaining cards to background batches
        self._pending_render = (visible, total_filtered, first_n,
                                self._render_generation)
        self._batch_after_id = self.after(1, self._render_next_batch)

    def _render_card_at(self, idx: int, oficio: Dict[str, Any]) -> None:
        card = self._card_pool[idx]
        card.update(oficio)
        col = idx % 2
        card.show(idx // 2, col, 0 if col == 0 else 6)

    def _render_next_batch(self) -> None:
        if self._pending_render is None:
            return
        visible, total_filtered, start, gen = self._pending_render
        if gen != self._render_generation or not self._card_pool:
            self._pending_render = None
            self._batch_after_id = None
            return

        BATCH = 4
        end = min(start + BATCH, len(visible))
        for idx in range(start, end):
            self._render_card_at(idx, visible[idx])

        if end >= len(visible):
            self._finalize_render(total_filtered, len(visible))
            self._pending_render = None
            self._batch_after_id = None
        else:
            self._pending_render = (visible, total_filtered, end, gen)
            self._batch_after_id = self.after(1, self._render_next_batch)

    def _finalize_render(self, total_filtered: int, visible_count: int) -> None:
        if self._more_widget is None:
            return
        if total_filtered > self._MAX_CARDS:
            more = total_filtered - self._MAX_CARDS
            self._more_widget.config(
                text=f"… y {more} oficio{'s' if more > 1 else ''} más — usa la búsqueda para refinar.",
            )
            self._more_widget.grid(
                row=(visible_count // 2) + 1, column=0, columnspan=2, pady=(8, 0),
            )
        else:
            self._more_widget.grid_forget()

# ---------------------------------------------------------------------------
# CLI entry point
# ---------------------------------------------------------------------------

def launch_main_gui(config: Config) -> None:
    ensure_excel_exists(config.excel_path)
    init_db(config.db_path)
    migrate_excel_to_db(config.excel_path, config.db_path)
    app = OficiosApp(config)
    app.mainloop()


def main() -> None:
    parser = argparse.ArgumentParser(
        description="Gestión de Oficios CGE — abre la interfaz gráfica si no se pasan argumentos."
    )
    default_config = get_base_dir() / "config.json"
    parser.add_argument("--config", default=str(default_config),
                        help="Ruta al archivo de configuración JSON.")
    parser.add_argument("--run-once", action="store_true",
                        help="Ejecuta una sola vez y termina (sin GUI).")
    parser.add_argument("--service", action="store_true",
                        help="Modo servicio continuo (sin GUI).")
    parser.add_argument("--create-template", action="store_true",
                        help="Crea solo la plantilla Excel y termina.")
    parser.add_argument("--reset", action="store_true",
                        help="Resetea la memoria de PDFs procesados.")
    parser.add_argument("--revaluar", action="store_true",
                        help="Abre interfaz de revaloración (legacy).")
    parser.add_argument("--stats", action="store_true",
                        help="Imprime métricas de accuracy del agente.")
    parser.add_argument("--regenerar-reglas", action="store_true",
                        help="Regenera reglas de clasificación con Claude.")
    args = parser.parse_args()

    config = load_config(Path(args.config))
    setup_logging(config.log_path)

    if args.create_template:
        create_excel_template(config.excel_path)
        logging.info("Plantilla creada: %s", config.excel_path)
        return

    if args.reset:
        reset_state(config)
        return

    if args.stats:
        historial = load_historial(config.historial_path)
        if not historial:
            print("No hay historial aún. Procesa algunos oficios primero.")
            return
        s = compute_learning_stats(historial)
        print(f"\nTotal decisiones: {s['total']}")
        print(f"Accuracy global:  {s['accuracy'] * 100:.1f}%  "
              f"({s['correctos']}/{s['total']})")
        if s["por_area"]:
            print("\nAccuracy por área:")
            for area, d in s["por_area"].items():
                bar = "█" * int(d["accuracy"] * 20) + "░" * (20 - int(d["accuracy"] * 20))
                print(f"  {area:<22} {bar}  {d['accuracy'] * 100:>3.0f}%  "
                      f"({d['correctos']}/{d['total']})")
        if s["errores_frecuentes"]:
            print("\nErrores más frecuentes:")
            for err, count in s["errores_frecuentes"]:
                print(f"  {err:<45} x{count}")
        return

    if args.regenerar_reglas:
        historial = load_historial(config.historial_path)
        print(f"Regenerando reglas con {len(historial)} entradas del historial...")
        reglas = generate_reglas(config, historial)
        if reglas:
            print(f"\nReglas guardadas en {config.reglas_path}\n")
            print(reglas)
        else:
            print("No se pudieron generar reglas (revisa logs).")
        return

    if args.run_once:
        run_once(config)
    elif args.service:
        service_loop(config)
    else:
        launch_main_gui(config)


if __name__ == "__main__":
    main()
